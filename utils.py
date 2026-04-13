import os
import asyncio
import logging
import base64
from google import genai
from google.genai import types
from PIL import Image, ImageEnhance
from dotenv import load_dotenv
from openai import AsyncOpenAI

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
openai_api_key = os.getenv("OPENAI_API_KEY")

# OpenAI Client (Asosiy tezkor)
openai_client = AsyncOpenAI(api_key=openai_api_key)

# GEMINI_BASE_URL faqat .env da bo'lsagina ishlatiladi.
base_url = os.getenv("GEMINI_BASE_URL")

if base_url and len(base_url.strip()) > 5:
    client = genai.Client(api_key=api_key, http_options={'base_url': base_url, 'timeout': 30.0})
else:
    client = genai.Client(api_key=api_key, http_options={'timeout': 30.0})

# Ovoz va boshqa og'ir zaproslar uchun to'g'ridan-to'g'ri API
direct_client = genai.Client(api_key=api_key, http_options={'timeout': 60.0})

CYRILLIC_TO_LATIN = {
    'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ё': 'Yo', 'Ж': 'J',
    'З': 'Z', 'И': 'I', 'Й': 'Y', 'К': 'K', 'Л': 'L', 'М': 'M', 'Н': 'N', 'О': 'O',
    'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U', 'Ф': 'F', 'Х': 'X', 'Ц': 'Ts',
    'Ч': 'Ch', 'Ш': 'Sh', 'Щ': 'Sh', 'Ъ': "'", 'Ы': 'I', 'Ь': '', 'Э': 'E', 'Ю': 'Yu',
    'Я': 'Ya', 'Ў': 'O\'', 'Қ': 'Q', 'Ғ': 'G\'', 'Ҳ': 'H',
    
    'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'yo', 'ж': 'j',
    'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm', 'н': 'n', 'о': 'o',
    'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u', 'ф': 'f', 'х': 'x', 'ц': 'ts',
    'ч': 'ch', 'ш': 'sh', 'щ': 'sh', 'ъ': "'", 'ы': 'i', 'ь': '', 'э': 'e', 'ю': 'yu',
    'я': 'ya', 'ў': "o'", 'қ': 'q', 'ғ': "g'", 'ҳ': 'h'
}

def transliterate_to_latin(text: str) -> str:
    """Kirill yozuvini Lotin yozuviga o'zgartiradi."""
    if not text:
        return text
    result = []
    i = 0
    while i < len(text):
        char = text[i]
        
        # Maxsus o'zgarishlar (misol: E harfi probeldan yoki unlidan keyin 'Ye' bo'ladi, boshqa payt 'E')
        if char == 'Е':
            if i == 0 or text[i-1] in ' \n\t-АЕЁИОУЫЭЮЯаеёиоуыэюя':
                result.append('Ye')
            else:
                result.append('e')
        elif char == 'е':
            if i == 0 or text[i-1] in ' \n\t-АЕЁИОУЫЭЮЯаеёиоуыэюя':
                result.append('ye')
            else:
                result.append('e')
        elif char == 'Ц':
            if i == 0 or text[i-1] in ' \n\t-АЕЁИОУЫЭЮЯаеёиоуыэюя':
                result.append('S')
            else:
                result.append('Ts')
        elif char == 'ц':
            if i == 0 or text[i-1] in ' \n\t-АЕЁИОУЫЭЮЯаеёиоуыэюя':
                result.append('s')
            else:
                result.append('ts')
        else:
            result.append(CYRILLIC_TO_LATIN.get(char, char))
        i += 1
    return "".join(result)

# ==================== OCR PROMPT ====================

STRUCTURED_OCR_PROMPT = """You are an expert document OCR system for official Uzbek documents.
Extract ALL text from this image completely and accurately. Do not skip any line or word.
This includes HANDWRITTEN text — carefully read and transcribe handwriting as well.

OUTPUT FORMAT RULES:
1. Each real paragraph = one continuous block of text (single long line). Do NOT break paragraphs into visual lines.
2. Separate paragraphs with exactly ONE empty line.
3. If a line is CENTERED (titles/headers), put [CENTER] before it on its own line.
4. If a paragraph has first-line indentation, put [INDENT] before it.
5. Do NOT use ** or any bold markers. Just output plain text.
6. Keep dash/bullet list items (- item) each on their own line with NO empty lines between consecutive dash items.
7. Do NOT add commentary, code blocks, or explanations.
8. Extract EVERY single word — do not skip or summarize anything.
9. FIX any obvious spelling and punctuation errors in the extracted text.

ALPHABET: Output MUST be in {alphabet}.
- 'Lotin' → O'zbek Latin
- 'Kirill' → O'zbek Cyrillic

Return ONLY the extracted text."""


# ==================== LOTIN <-> KIRILL KONVERTOR ====================

# O'zbek lotin -> kirill xaritasi
_LATIN_TO_CYRILLIC = {
    "yo'": "йў", "Yu": "Ю", "yu": "ю", "Ya": "Я", "ya": "я",
    "Yo": "Ё", "yo": "ё", "Ye": "Е", "ye": "е",
    "Sh": "Ш", "sh": "ш", "Ch": "Ч", "ch": "ч",
    "G'": "Ғ", "g'": "ғ", "O'": "Ў", "o'": "ў",
    "Ng": "Нг", "ng": "нг", "Ts": "Ц", "ts": "ц",
    "A": "А", "a": "а", "B": "Б", "b": "б",
    "D": "Д", "d": "д", "E": "Э", "e": "э",
    "F": "Ф", "f": "ф", "G": "Г", "g": "г",
    "H": "Ҳ", "h": "ҳ", "I": "И", "i": "и",
    "J": "Ж", "j": "ж", "K": "К", "k": "к",
    "L": "Л", "l": "л", "M": "М", "m": "м",
    "N": "Н", "n": "н", "O": "О", "o": "о",
    "P": "П", "p": "п", "Q": "Қ", "q": "қ",
    "R": "Р", "r": "р", "S": "С", "s": "с",
    "T": "Т", "t": "т", "U": "У", "u": "у",
    "V": "В", "v": "в", "X": "Х", "x": "х",
    "Y": "Й", "y": "й", "Z": "З", "z": "з",
    "'": "ъ",
}

# Kirill -> lotin xaritasi
_CYRILLIC_TO_LATIN = {
    "Ё": "Yo", "ё": "yo", "Ю": "Yu", "ю": "yu",
    "Я": "Ya", "я": "ya", "Ш": "Sh", "ш": "sh",
    "Ч": "Ch", "ч": "ch", "Ғ": "G'", "ғ": "g'",
    "Ў": "O'", "ў": "o'", "Ц": "Ts", "ц": "ts",
    "Щ": "Sh", "щ": "sh", "Ъ": "'", "ъ": "'",
    "Ь": "", "ь": "",
    "А": "A", "а": "a", "Б": "B", "б": "b",
    "В": "V", "в": "v", "Г": "G", "г": "g",
    "Д": "D", "д": "d", "Е": "E", "е": "e",
    "Ж": "J", "ж": "j", "З": "Z", "з": "z",
    "И": "I", "и": "i", "Й": "Y", "й": "y",
    "К": "K", "к": "k", "Қ": "Q", "қ": "q",
    "Л": "L", "л": "l", "М": "M", "м": "m",
    "Н": "N", "н": "n", "О": "O", "о": "o",
    "П": "P", "п": "p", "Р": "R", "р": "r",
    "С": "S", "с": "s", "Т": "T", "т": "t",
    "У": "U", "у": "u", "Ф": "F", "ф": "f",
    "Х": "X", "х": "x", "Ҳ": "H", "ҳ": "h",
}


def convert_latin_to_cyrillic(text: str) -> str:
    """Lotin yozuvini Kirill yozuviga o'girish."""
    result = text
    for lat, cyr in sorted(_LATIN_TO_CYRILLIC.items(), key=lambda x: -len(x[0])):
        result = result.replace(lat, cyr)
    return result


def convert_cyrillic_to_latin(text: str) -> str:
    """Kirill yozuvini Lotin yozuviga o'girish."""
    result = text
    for cyr, lat in sorted(_CYRILLIC_TO_LATIN.items(), key=lambda x: -len(x[0])):
        result = result.replace(cyr, lat)
    return result


def detect_script(text: str) -> str:
    """Matnning asosiy yozuvi qaysi ekanligini aniqlash."""
    cyrillic_count = 0
    latin_count = 0
    for ch in text:
        if '\u0400' <= ch <= '\u04FF':
            cyrillic_count += 1
        elif 'a' <= ch.lower() <= 'z':
            latin_count += 1
    if cyrillic_count > latin_count:
        return "kirill"
    return "lotin"


def process_docx_alphabet(docx_path: str, alphabet: str):
    """Word fayli matnini xavfsiz (docx library orqali) o'zgartiradi."""
    import logging
    try:
        import docx
        doc = docx.Document(docx_path)
        
        def process_text(text: str) -> str:
            if not text.strip():
                return text
            if alphabet == "Kirill":
                from utils import convert_latin_to_cyrillic
                return convert_latin_to_cyrillic(text)
            else:
                from utils import convert_cyrillic_to_latin
                return convert_cyrillic_to_latin(text)
        
        # Paragraflarni tahlil qilish
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text:
                    run.text = process_text(run.text)
                    
        # Jadvallarni tahlil qilish
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if run.text:
                                run.text = process_text(run.text)
                                
        doc.save(docx_path)
        logging.info(f"process_docx_alphabet (python-docx) orqali yozuvlar muvaffaqiyatli almashtirildi: {docx_path}")
    except Exception as e:
        logging.error(f"process_docx_alphabet xatosi (Fayl o'zgartirilmadi): {e}")


# ==================== RASM YAXSHILASH ====================

def _enhance_image(img: Image.Image) -> Image.Image:
    """Rasm sifatini avtomatik yaxshilash — OCR uchun optimallash."""
    try:
        # Kontrast oshirish
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.4)
        
        # Aniqlik (sharpness) oshirish
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(1.6)
        
        # Yorug'lik biroz oshirish
        enhancer = ImageEnhance.Brightness(img)
        img = enhancer.enhance(1.1)
    except Exception as e:
        logging.warning(f"Rasm yaxshilashda xato: {e}")
    
    return img


# ==================== RASM QAYTA ISHLASH ====================

async def process_image_async(image_path: str, alphabet: str) -> tuple:
    import json
    import base64
    image = await asyncio.to_thread(_open_and_optimize_image, image_path)
    
    prompt = (
        f"Analyze this document image. We are reconstructing the layout automatically. "
        f"1. Extract all text completely and translate/convert it carefully into the {alphabet} alphabet. "
        f"2. Identify the `text_bottom_y` (0-1000 scale) coordinate where the LAST line of the text block above the diagram ends. "
        f"3. Identify the bounding boxes of ONLY the non-text pure graphical elements (e.g. the exact outline of the engine or device). "
        f"CRITICAL: The graphical diagram's bounding box `ymin` MUST be strictly greater than `text_bottom_y` to ensure no text is included in the bounding box. "
        f"Return ONLY valid JSON exactly like this: "
        f"{{\n"
        f"  \"text\": \"full extracted and converted text\",\n"
        f"  \"text_bottom_y\": 250,\n"
        f"  \"diagrams\": [ {{\"box_2d\": [ymin, xmin, ymax, xmax]}} ]\n"
        f"}}\n"
        f"If there are no physical diagrams, leave the diagrams array empty."
    )

    try:
        import io
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format='JPEG')
        base64_image = base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')
        
        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                                "detail": "high"
                            }
                        }
                    ]
                }
            ],
            temperature=0.1
        )
        
        if response.choices and response.choices[0].message.content:
            try:
                raw_json = response.choices[0].message.content.strip()
                data = json.loads(raw_json)
                text_result = data.get("text", "").strip()
                diagrams = data.get("diagrams", [])
                
                cropped_path = None
                if diagrams and len(diagrams) > 0:
                    box = None
                    diagram_item = diagrams[0]
                    if isinstance(diagram_item, dict):
                        box = diagram_item.get("box_2d", diagram_item.get("box_3d"))
                    elif isinstance(diagram_item, list):
                        box = diagram_item
                    
                    if isinstance(box, list) and len(box) == 4:
                        try:
                            ymin, xmin, ymax, xmax = [int(v) for v in box]
                            text_bottom = int(data.get("text_bottom_y", 0))
                            if text_bottom > 0 and text_bottom < ymax:
                                ymin = max(ymin, text_bottom + 5) 
                            
                            ymin = min(ymin + 65, ymax - 30)
                            ymax = max(ymax - 25, ymin + 20)
                            xmin = min(xmin + 25, xmax - 10)
                            xmax = max(xmax - 20, xmin + 10)
                            
                            w, h = image.size
                            left = int(xmin * w / 1000)
                            top = int(ymin * h / 1000)
                            right = int(xmax * w / 1000)
                            bottom = int(ymax * h / 1000)
                            
                            if right > left and bottom > top:
                                cropped = image.crop((left, top, right, bottom))
                                cropped_path = f"{image_path}_crop.jpg"
                                if cropped.mode in ("RGBA", "P"):
                                    cropped = cropped.convert("RGB")
                                cropped.save(cropped_path, "JPEG", quality=95)
                        except Exception as parse_err:
                            logging.error(f"Kordinata xatoligi: {parse_err}")
                
                return text_result, cropped_path
                
            except Exception as e:
                logging.error(f"JSON parsing error yoki rasm kesishda xato: {e}")
                return raw_json[:3900], None
        else:
            raise ValueError("OpenAI API bo'sh javob qaytardi.")
    except Exception as e:
        logging.error(f"OpenAI API xatolik: {e}")
        raise


def _open_and_optimize_image(image_path: str) -> Image.Image:
    img = Image.open(image_path)
    max_dimension = 2048
    if img.width > max_dimension or img.height > max_dimension:
        img.thumbnail((max_dimension, max_dimension), Image.Resampling.LANCZOS)
    if img.mode == 'RGBA':
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg
    elif img.mode != 'RGB':
        img = img.convert('RGB')
    
    # Avtomatik rasm yaxshilash
    img = _enhance_image(img)
    
    return img


# ==================== MATN TARJIMASI ====================

async def translate_text(text: str, target_lang: str) -> str:
    """Matnni boshqa tilga tarjima qilish."""
    lang_map = {
        "uz_ru": "O'zbek tilidan Rus tiliga",
        "ru_uz": "Rus tilidan O'zbek tiliga", 
        "uz_en": "O'zbek tilidan Ingliz tiliga",
        "en_uz": "Ingliz tilidan O'zbek tiliga",
    }
    direction = lang_map.get(target_lang, "boshqa tilga")
    
    prompt = (
        f"{direction} tarjima qiling. Mazmunini o'zgartirmang.\n"
        f"Faqat tarjima qilingan matnni qaytaring, boshqa hech narsa qo'shmang.\n\n"
        f"Matn:\n{text}"
    )
    
    try:
        response = await client.aio.models.generate_content(
            model='gemini-2.5-flash',
            contents=[prompt]
        )
        if response.text:
            return response.text.strip()
        return text
    except Exception as e:
        logging.error(f"Tarjima xatolik: {e}")
        return text


# ==================== OVOZLI XABAR ====================

async def process_voice_async(voice_path: str, alphabet: str) -> str:
    """Ovozli xabarni matnga aylantirish. OpenAI Whisper."""
    import logging
    try:
        logging.info("Ovoz: OpenAI Whisper usuli sinab ko'rilmoqda...")
        with open(voice_path, 'rb') as audio_file:
            response = await openai_client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                prompt="O'zbek tilidagi jonli ovozli xabar. Iltimos, so'zlarni xatosiz va tushunarli qilib matnga o'giring."
            )
        
        text = response.text
        if text and len(text.strip()) > 1:
            logging.info(f"Ovoz: OpenAI Whisper muvaffaqiyatli! ({len(text)} belgi)")
            # Alphabetni tekshirib to'g'rilash (Whisper ba'zan Kirill yoki Lotinning aralashida berishi mumkin)
            if alphabet == "Kirill":
                from utils import convert_latin_to_cyrillic
                text = convert_latin_to_cyrillic(text)
            elif alphabet == "Lotin":
                from utils import transliterate_to_latin
                text = transliterate_to_latin(text)
            return text.strip()
            
    except Exception as e:
        logging.warning(f"Ovoz: OpenAI Whisper xato berdi: {e}")
        
    raise Exception("Ovoz tahlili muvaffaqiyatsiz. Iltimos, aniqroq gapirgan holda qayta yuboring.")


def _speech_recognition_sync(voice_path: str) -> str:
    """SpeechRecognition bilan sinxron ovoz tahlili."""
    import subprocess
    import sys
    import tempfile
    
    # speech_recognition va pydub o'rnatish (agar yo'q bo'lsa)
    try:
        import speech_recognition as sr
    except ImportError:
        subprocess.run([sys.executable, "-m", "pip", "install", "SpeechRecognition", "pydub"], 
                      capture_output=True, timeout=60)
        import speech_recognition as sr
    
    # OGG -> WAV konvertatsiya (ffmpeg yoki pydub orqali)
    wav_path = voice_path + ".wav"
    ffmpeg_exe = "ffmpeg"
    
    import shutil
    import urllib.request
    import zipfile
    
    # ffmpeg o'rnatilganini tekshiramiz
    if not shutil.which("ffmpeg"):
        ffmpeg_bin_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ffmpeg-bin")
        local_ffmpeg = os.path.join(ffmpeg_bin_dir, "ffmpeg.exe")
        
        if not os.path.exists(local_ffmpeg):
            try:
                logging.info("ffmpep topilmadi! Standalone versiya yuklanmoqda...")
                os.makedirs(ffmpeg_bin_dir, exist_ok=True)
                zip_path = os.path.join(ffmpeg_bin_dir, "ffmpeg.zip")
                urllib.request.urlretrieve("https://github.com/BtbN/FFmpeg-Builds/releases/download/latest/ffmpeg-master-latest-win64-gpl.zip", zip_path)
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    for el in zip_ref.namelist():
                        if el.endswith("ffmpeg.exe"):
                            zip_ref.extract(el, ffmpeg_bin_dir)
                            extracted_path = os.path.join(ffmpeg_bin_dir, el)
                            shutil.move(extracted_path, local_ffmpeg)
                            break
                os.remove(zip_path)
            except Exception as e:
                logging.warning(f"ffmpeg yuklashda xatolik: {e}")
        
        if os.path.exists(local_ffmpeg):
            ffmpeg_exe = local_ffmpeg
    
    try:
        # ffmpeg orqali konvertatsiya (eng ishonchli)
        result = subprocess.run(
            [ffmpeg_exe, "-i", voice_path, "-ar", "16000", "-ac", "1", "-y", wav_path],
            capture_output=True, timeout=30
        )
        if result.returncode != 0:
            raise Exception("ffmpeg ishlamadi")
    except (FileNotFoundError, Exception):
        # pydub orqali
        try:
            from pydub import AudioSegment
            # Pydub uchun ham ffmpeg yo'lini ko'rsatamiz
            if ffmpeg_exe != "ffmpeg":
                AudioSegment.converter = ffmpeg_exe
            audio = AudioSegment.from_ogg(voice_path)
            audio = audio.set_frame_rate(16000).set_channels(1)
            audio.export(wav_path, format="wav")
        except Exception as e:
            raise Exception(f"Ovoz faylini WAV ga o'tkazib bo'lmadi: {e}")
    
    # Speech Recognition
    if not os.path.exists(wav_path):
        raise Exception("WAV fayl yaratilmadi!")
        
    recognizer = sr.Recognizer()
    with sr.AudioFile(wav_path) as source:
        audio = recognizer.record(source)
    
    # Tozalash
    try:
        os.remove(wav_path)
    except:
        pass
    
    # Google Speech Recognition (bepul, 50 ta so'rovgacha)
    try:
        text = recognizer.recognize_google(audio, language="uz-UZ")
        return text
    except sr.UnknownValueError:
        # O'zbek tilida topolmasa ruscha sinab ko'rish
        try:
            text = recognizer.recognize_google(audio, language="ru-RU")
            return text
        except:
            pass
    except Exception:
        pass
    
    # Oxirgi iloj - inglizcha
    try:
        text = recognizer.recognize_google(audio, language="en-US")
        return text
    except Exception as e:
        raise Exception(f"Google Speech API ham ishlamadi: {e}")


# ==================== SIFAT BALLI ====================

async def calculate_quality_score(text: str) -> dict:
    """Hujjat sifatini AI bilan baholash."""
    prompt = (
        "Quyidagi matnning sifatini 3 ta mezon bo'yicha 0-100 ball bilan baholang:\n"
        "1. OCR aniqligi (imlo, to'liqlik)\n"
        "2. Formatlash sifati (paragraflar, tuzilish)\n"
        "3. Grammatika va tinish belgilari\n\n"
        "Faqat 3 ta sonni qaytaring, vergul bilan ajratilgan. Masalan: 95,88,92\n"
        "Boshqa hech narsa yozmang.\n\n"
        f"Matn:\n{text[:2000]}"
    )
    
    try:
        response = await client.aio.models.generate_content(
            model='gemini-2.5-flash',
            contents=[prompt]
        )
        if response.text:
            scores = response.text.strip().replace(" ", "").split(",")
            if len(scores) >= 3:
                ocr = min(100, max(0, int(scores[0])))
                formatting = min(100, max(0, int(scores[1])))
                grammar = min(100, max(0, int(scores[2])))
                overall = round((ocr + formatting + grammar) / 3)
                return {
                    "ocr": ocr,
                    "formatting": formatting,
                    "grammar": grammar,
                    "overall": overall,
                }
    except Exception as e:
        logging.warning(f"Sifat balli xatolik: {e}")
    
    return {"ocr": 0, "formatting": 0, "grammar": 0, "overall": 0}

# os va fitz allaqachon yuqorida import qilingan
# pytesseract va Document faqat kerak bo'lganda import qilinadi (lazy)


def extract_text_via_tesseract(pdf_path: str, docx_path: str, alphabet: str):
    """
    Skanerlangan PDF fayldan rasmlarni kesib olib, Tesseract yordamida o'qiydi 
    va transliteratsiya qilib Word (docx) fayliga joylaydi.
    """
    import fitz
    import pytesseract
    from docx import Document
    
    # Tesseract executable path
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    
    doc = fitz.open(pdf_path)
    word_doc = Document()
    
    # Tillarni sozlash
    tess_lang = 'rus+uzb_cyrl+eng'  # Default for Kirill
    if alphabet == 'Lotin':
        tess_lang = 'uzb+eng'
        
    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=300)
        temp_img = f"temp_tess_page_{page_num}.png"
        pix.save(temp_img)
        
        try:
            # Rasm ustida OCR
            text = pytesseract.image_to_string(temp_img, lang=tess_lang)
            text = text.strip()
            
            # Agar text topilmasa, tashlab o'tamiz
            if text:
                # Transliteratsiya
                if alphabet == 'Lotin':
                    text = convert_cyrillic_to_latin(text)
                elif alphabet == 'Kirill':
                    text = convert_latin_to_cyrillic(text)
                
                # Word ga qo'shish
                for para in text.split('\n'):
                    if para.strip():
                        word_doc.add_paragraph(para.strip())
                
                # Yangi sahifa
                if page_num < len(doc) - 1:
                    word_doc.add_page_break()
                    
        except Exception as e:
            logging.error(f"OCR Error on page {page_num}: {e}")
        finally:
            if os.path.exists(temp_img):
                os.remove(temp_img)
                
    # Faylni saqlash
    word_doc.save(docx_path)
    doc.close()


def convert_pdf_to_word_formatted(pdf_path: str, docx_path: str, alphabet: str = "Original"):
    """
    Professional PDF → Word konvertor.
    
    Asosiy xususiyatlar:
    - SATR darajasida markazlash aniqlash (blok emas!)
    - Body matnni birlashtirish (har satr alohida paragraf bo'lmaydi)
    - Bold/italic/shrift o'lchami saqlash
    - Tire ro'yxat formatlash
    - Abzats boshi (indent) aniqlash
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import fitz
    
    pdf_doc = fitz.open(pdf_path)
    word_doc = Document()
    
    # Word hujjat sozlamalari
    for section in word_doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    logging.info(f"PDF→Word formatted v2: {len(pdf_doc)} sahifa, alphabet={alphabet}")
    
    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        page_width = page.rect.width
        page_center = page_width / 2
        
        # Sahifaning haqiqiy matn chegaralarini aniqlash
        # Barcha satrlarning chap va o'ng chegaralarini yig'amiz
        all_lines_info = []
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        text_blocks = [b for b in blocks if b["type"] == 0]
        text_blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
        
        # 1-qadam: Barcha satrlarni tahlil qilish
        for block in text_blocks:
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                
                # Satr matnini yig'ish
                line_text = ""
                for sp in spans:
                    line_text += sp.get("text", "")
                line_text_stripped = line_text.strip()
                if not line_text_stripped:
                    continue
                
                # Satrning haqiqiy chap va o'ng chegarasini span lardan olish
                line_x0 = min(sp["bbox"][0] for sp in spans if sp.get("text", "").strip())
                line_x1 = max(sp["bbox"][2] for sp in spans if sp.get("text", "").strip())
                line_width = line_x1 - line_x0
                line_center = (line_x0 + line_x1) / 2
                line_y = line["bbox"][1]
                
                # Span ma'lumotlarini saqlash
                all_lines_info.append({
                    "spans": spans,
                    "text": line_text_stripped,
                    "x0": line_x0,
                    "x1": line_x1,
                    "width": line_width,
                    "center": line_center,
                    "y": line_y,
                    "y1": line["bbox"][3],
                    "block_num": block["number"] if "number" in block else id(block),
                })
        
        if not all_lines_info:
            continue
        
        # Sahifaning matn chegaralarini aniqlash (eng keng satrlardan)
        if len(all_lines_info) > 2:
            widths = sorted([li["width"] for li in all_lines_info], reverse=True)
            max_text_width = widths[min(2, len(widths)-1)]  # 3-eng keng satr
        else:
            max_text_width = max(li["width"] for li in all_lines_info)
        
        # 2-qadam: Har bir satr uchun format aniqlash va paragraflar yaratish
        i = 0
        while i < len(all_lines_info):
            li = all_lines_info[i]
            
            # --- MARKAZLASHNI ANIQLASH (satr darajasida) ---
            # Satr markazi sahifa markaziga yaqin va satr kengligi to'liq kenglikdan kichik
            center_diff = abs(li["center"] - page_center)
            is_line_centered = (
                center_diff < page_width * 0.08  # markazdan ±8% ichida
                and li["width"] < max_text_width * 0.92  # to'liq kenglikdan kichik
            )
            
            # --- TIRE RO'YXAT ---
            is_dash = li["text"].startswith("- ") or li["text"].startswith("— ") or li["text"].startswith("– ")
            
            if is_line_centered:
                # Markazlangan satr — alohida paragraf
                para = word_doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = 1.15
                
                _add_formatted_spans(para, li["spans"], alphabet)
                i += 1
                
            elif is_dash:
                # Tire ro'yxat elementi
                para = word_doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.left_indent = Cm(1.0)
                para.paragraph_format.first_line_indent = Cm(-0.5)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = 1.15
                
                _add_formatted_spans(para, li["spans"], alphabet)
                i += 1
                
            else:
                # Oddiy matn — ketma-ket satrlarni BIR paragrafga birlashtirish
                para = word_doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.15
                para.paragraph_format.first_line_indent = Cm(1.25)
                
                _add_formatted_spans(para, li["spans"], alphabet)
                
                # Keyingi satrlarni ham shu paragrafga qo'shish
                # (agar ular ham oddiy matn bo'lsa va bir xil blokda bo'lsa)
                j = i + 1
                while j < len(all_lines_info):
                    next_li = all_lines_info[j]
                    next_center_diff = abs(next_li["center"] - page_center)
                    next_is_centered = (
                        next_center_diff < page_width * 0.08
                        and next_li["width"] < max_text_width * 0.92
                    )
                    next_is_dash = (
                        next_li["text"].startswith("- ") or 
                        next_li["text"].startswith("— ") or 
                        next_li["text"].startswith("– ")
                    )
                    
                    # Agar keyingi satr markazlangan yoki tire yoki boshqa blokda — to'xtatish
                    if next_is_centered or next_is_dash:
                        break
                    
                    # Agar satrlar orasida katta bo'shliq bo'lsa — yangi paragraf
                    vertical_gap = next_li["y"] - li["y1"]
                    line_height = li["y1"] - li["y"]
                    if line_height > 0 and vertical_gap > line_height * 1.5:
                        break
                    
                    # Agar boshqa blokda va katta bo'shliq — yangi paragraf
                    if next_li["block_num"] != li["block_num"]:
                        if vertical_gap > line_height * 0.8:
                            break
                    
                    # Bu satrni paragrafga qo'shish (probel bilan)
                    run = para.add_run(" ")
                    _add_formatted_spans(para, next_li["spans"], alphabet)
                    
                    li = next_li
                    j += 1
                
                i = j
        
        # Sahifalar orasida page break
        if page_num < len(pdf_doc) - 1:
            word_doc.add_page_break()
    
    # Word faylni saqlash
    word_doc.save(docx_path)
    pdf_doc.close()
    logging.info(f"PDF→Word formatted v2 muvaffaqiyatli: {docx_path}")


def _add_formatted_spans(para, spans, alphabet):
    """Span larni paragrafga formatlangan holda qo'shish."""
    for span in spans:
        text = span.get("text", "")
        if not text:
            continue
        
        font_name = span.get("font", "")
        font_size = span.get("size", 12)
        font_flags = span.get("flags", 0)
        
        # Bold aniqlash: flags bit 4 YOKI font nomida Bold/bold/BOLD
        is_bold = (
            bool(font_flags & 16) or 
            "Bold" in font_name or 
            "bold" in font_name or 
            "BOLD" in font_name or
            "Black" in font_name
        )
        # Italic aniqlash
        is_italic = (
            bool(font_flags & 2) or 
            "Italic" in font_name or 
            "italic" in font_name or
            "Oblique" in font_name
        )
        
        # Transliteratsiya
        if alphabet == "Lotin":
            text = convert_cyrillic_to_latin(text)
        elif alphabet == "Kirill":
            text = convert_latin_to_cyrillic(text)
        
        # Run yaratish
        run = para.add_run(text)
        run.font.size = Pt(max(8, min(28, font_size)))
        run.bold = is_bold
        run.italic = is_italic
        
        # Shrift nomi
        if any(kw in font_name.lower() for kw in ["times", "serif"]):
            run.font.name = "Times New Roman"
        elif any(kw in font_name.lower() for kw in ["arial", "helvetica", "sans"]):
            run.font.name = "Arial"
        else:
            run.font.name = "Times New Roman"


