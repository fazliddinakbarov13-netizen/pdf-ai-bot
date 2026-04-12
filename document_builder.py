import os
import logging
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


BOT_WATERMARK = ""


# ==================== PARAGRAF TAHLILI ====================

def parse_paragraphs(text: str) -> list:
    raw_lines = text.split("\n")
    result = []
    
    for line in raw_lines:
        stripped = line.strip()
        
        if not stripped:
            result.append({"type": "empty"})
            continue
        
        centered = False
        indent = False
        content = stripped
        
        if content.startswith("[CENTER]"):
            centered = True
            content = content.replace("[CENTER]", "", 1).strip()
        
        if content.startswith("[INDENT]"):
            indent = True
            content = content.replace("[INDENT]", "", 1).strip()
        
        content = content.replace("**", "")
        
        is_dash = content.startswith(("- ", "– ", "• "))
        
        result.append({
            "type": "text",
            "content": content,
            "centered": centered,
            "indent": indent,
            "dash": is_dash,
        })
    
    return result


def optimize_blocks(blocks: list) -> list:
    """Ketma-ket dash qatorlari orasidagi bo'sh qatorlarni olib tashlash."""
    optimized = []
    i = 0
    while i < len(blocks):
        block = blocks[i]
        
        if block["type"] == "empty":
            prev_dash = (len(optimized) > 0 and 
                        optimized[-1]["type"] == "text" and 
                        optimized[-1].get("dash", False))
            next_idx = i + 1
            while next_idx < len(blocks) and blocks[next_idx]["type"] == "empty":
                next_idx += 1
            next_dash = (next_idx < len(blocks) and 
                        blocks[next_idx]["type"] == "text" and 
                        blocks[next_idx].get("dash", False))
            
            if prev_dash and next_dash:
                i += 1
                continue
        
        optimized.append(block)
        i += 1
    
    return optimized


# ==================== WORD ====================

def create_word_document(text: str, output_path: str, image_path: str = None):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    blocks = optimize_blocks(parse_paragraphs(text))
    
    prev_was_dash = False
    
    for i, block in enumerate(blocks):
        if block["type"] == "empty":
            next_is_dash = (i + 1 < len(blocks) and 
                           blocks[i+1]["type"] == "text" and 
                           blocks[i+1].get("dash", False))
            if prev_was_dash and next_is_dash:
                continue
            prev_was_dash = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run("")
            run.font.size = Pt(2)
            continue
        
        content = block["content"]
        prev_was_dash = block.get("dash", False)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        if block["centered"]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if block["indent"]:
            p.paragraph_format.first_line_indent = Cm(1.25)
        
        run = p.add_run(content)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
    # Qirqib olingan vizual rasm kiritish (agar berilgan bo'lsa) yozuvdan KEYIN
    if image_path and os.path.exists(image_path):
        from PIL import Image
        try:
            img = Image.open(image_path)
            img_w_px, img_h_px = img.size
            img.close()
            max_w_cm, max_h_cm = 16.5, 12.0
            img_aspect = img_w_px / img_h_px
            if (max_w_cm / max_h_cm) <= img_aspect:
                display_w, display_h = max_w_cm, max_w_cm / img_aspect
            else:
                display_h, display_w = max_h_cm, max_h_cm * img_aspect
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(image_path, width=Cm(display_w), height=Cm(display_h))
        except Exception:
            pass
    
    # Watermark — hujjat pastidan
    if BOT_WATERMARK:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(BOT_WATERMARK)
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(180, 180, 180)
    
    doc.save(output_path)
    return output_path


# ==================== KO'P RASMLI WORD ====================

def create_multi_image_word(texts: list, output_path: str, image_paths: list = None):
    """Bir nechta rasm matnini bitta Word hujjatga yig'ish."""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    for idx, text in enumerate(texts):
        if idx > 0:
            doc.add_page_break()
            
        blocks = optimize_blocks(parse_paragraphs(text))
        prev_was_dash = False
        
        for i, block in enumerate(blocks):
            if block["type"] == "empty":
                next_is_dash = (i + 1 < len(blocks) and 
                               blocks[i+1]["type"] == "text" and 
                               blocks[i+1].get("dash", False))
                if prev_was_dash and next_is_dash:
                    continue
                prev_was_dash = False
                p = doc.add_paragraph()
                run = p.add_run("")
                run.font.size = Pt(2)
                continue
            
            content = block["content"]
            prev_was_dash = block.get("dash", False)
            
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            if block["centered"]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if block["indent"]:
                p.paragraph_format.first_line_indent = Cm(1.25)
            
            run = p.add_run(content)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            
        if image_paths and idx < len(image_paths) and image_paths[idx] and os.path.exists(image_paths[idx]):
            from PIL import Image
            try:
                img = Image.open(image_paths[idx])
                img_w_px, img_h_px = img.size
                img.close()
                max_w_cm, max_h_cm = 16.5, 12.0
                img_aspect = img_w_px / img_h_px
                if (max_w_cm / max_h_cm) <= img_aspect:
                    display_w, display_h = max_w_cm, max_w_cm / img_aspect
                else:
                    display_h, display_w = max_h_cm, max_h_cm * img_aspect
                
                doc.add_paragraph()  # bitta bo'sh joy
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(image_paths[idx], width=Cm(display_w), height=Cm(display_h))
            except Exception:
                pass
    
    # Watermark
    if BOT_WATERMARK:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(BOT_WATERMARK)
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(180, 180, 180)
    
    doc.save(output_path)
    return output_path


# ==================== PDF ====================

class CompactPDF(FPDF):
    def __init__(self):
        super().__init__()
        self._font_loaded = False
        self._load_font()
        self.set_margins(20, 15, 15)
        self.set_auto_page_break(auto=True, margin=20)
        self._total_pages = 0
    
    def _load_font(self):
        candidates = [
            # Windows
            (r"C:\Windows\Fonts\times.ttf", r"C:\Windows\Fonts\timesbd.ttf"),
            (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
            (r"C:\Windows\Fonts\calibri.ttf", r"C:\Windows\Fonts\calibrib.ttf"),
            # Linux
            ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
        ]
        for regular, bold in candidates:
            if os.path.exists(regular):
                try:
                    self.add_font("DocFont", "", regular)
                    if os.path.exists(bold):
                        self.add_font("DocFont", "B", bold)
                    self._font_loaded = True
                    break
                except Exception as e:
                    logging.warning(f"Shrift yuklanmadi: {e}")
    
    def set_doc_font(self, size=14, bold=False):
        if self._font_loaded:
            self.set_font("DocFont", "B" if bold else "", size)
        else:
            self.set_font("Arial", "B" if bold else "", size)
    
    def header(self):
        pass
    
    def footer(self):
        """Sahifa raqami va watermark pastdan."""
        self.set_y(-15)
        if self._font_loaded:
            self.set_font("DocFont", "", 8)
        else:
            self.set_font("Arial", "", 8)
        self.set_text_color(180, 180, 180)
        page_w = self.w - self.l_margin - self.r_margin
        if BOT_WATERMARK:
            self.cell(page_w / 2, 5, BOT_WATERMARK, 0, 0, 'L')
            self.cell(page_w / 2, 5, f"{self.page_no()}", 0, 0, 'R')
        else:
            self.cell(page_w, 5, f"{self.page_no()}", 0, 0, 'R')


def _sanitize(text):
    for old, new in {
        '\u200b': '', '\u200c': '', '\u200d': '', '\ufeff': '',
        '\u00a0': ' ', '\u2018': "'", '\u2019': "'",
        '\u201c': '"', '\u201d': '"', '\u2013': '-', '\u2014': '-',
        '\u2026': '...', '\u00ab': '"', '\u00bb': '"',
    }.items():
        text = text.replace(old, new)
    return text


def _render_text_to_pdf(pdf, text):
    """Matnni PDF ga yozish."""
    blocks = optimize_blocks(parse_paragraphs(text))
    lh = 6.0
    
    prev_was_dash = False
    
    for i, block in enumerate(blocks):
        if block["type"] == "empty":
            next_is_dash = (i + 1 < len(blocks) and 
                           blocks[i+1]["type"] == "text" and 
                           blocks[i+1].get("dash", False))
            if prev_was_dash and next_is_dash:
                continue
            prev_was_dash = False
            pdf.ln(1)
            continue
        
        content = block["content"]
        prev_was_dash = block.get("dash", False)
        
        if block["indent"]:
            content = "    " + content
        
        try:
            pdf.set_doc_font(14)
            pdf.set_text_color(0, 0, 0)
            
            if block["centered"]:
                pdf.multi_cell(0, lh, content, align='C')
            else:
                pdf.multi_cell(0, lh, content, align='L')
        except Exception as e:
            logging.warning(f"PDF xato: {e}")
            try:
                pdf.multi_cell(0, lh, content, align='L')
            except Exception:
                pdf.ln(lh)


def create_pdf_document(text: str, output_path: str, image_path: str = None):
    pdf = CompactPDF()
    pdf.add_page()
    text = _sanitize(text)
    _render_text_to_pdf(pdf, text)
    
    if image_path and os.path.exists(image_path):
        from PIL import Image
        try:
            img = Image.open(image_path)
            img_w_px, img_h_px = img.size
            img.close()
            page_w, page_h = 210, 297
            margin = 20
            usable_w = page_w - 2 * margin
            usable_h = 100
            img_aspect = img_w_px / img_h_px
            page_aspect = usable_w / usable_h
            if img_aspect >= page_aspect:
                display_w = usable_w
                display_h = display_w / img_aspect
            else:
                display_h = usable_h
                display_w = display_h * img_aspect
            
            pdf.ln(5)
            x = (page_w - display_w) / 2
            y = pdf.get_y()
            if y + display_h > page_h - margin:
                pdf.add_page()
                y = pdf.get_y()
            pdf.image(image_path, x=x, y=y, w=display_w, h=display_h)
        except Exception:
            pass
    pdf.output(output_path)
    return output_path


# ==================== KO'P RASMLI PDF ====================

def create_multi_image_pdf(texts: list, output_path: str, image_paths: list = None):
    """Bir nechta rasm matnini bitta PDF ga yig'ish."""
    pdf = CompactPDF()
    
    for idx, text in enumerate(texts):
        pdf.add_page()
        text = _sanitize(text)
        _render_text_to_pdf(pdf, text)
        
        if image_paths and idx < len(image_paths) and image_paths[idx] and os.path.exists(image_paths[idx]):
            from PIL import Image
            try:
                img = Image.open(image_paths[idx])
                img_w_px, img_h_px = img.size
                img.close()
                page_w, page_h = 210, 297
                margin = 20
                usable_w = page_w - 2 * margin
                usable_h = 100
                img_aspect = img_w_px / img_h_px
                page_aspect = usable_w / usable_h
                if img_aspect >= page_aspect:
                    display_w = usable_w
                    display_h = display_w / img_aspect
                else:
                    display_h = usable_h
                    display_w = display_h * img_aspect
                
                pdf.ln(5)
                x = (page_w - display_w) / 2
                y = pdf.get_y()
                if y + display_h > page_h - margin:
                    pdf.add_page()
                    y = pdf.get_y()
                pdf.image(image_paths[idx], x=x, y=y, w=display_w, h=display_h)
            except Exception:
                pass
    
    pdf.output(output_path)
    return output_path


# ==================== PAROLLI PDF ====================

def _add_password_to_pdf(input_path: str, output_path: str, password: str):
    """PDF ga parol qo'yish."""
    try:
        from PyPDF2 import PdfReader, PdfWriter
        
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page)
        
        writer.encrypt(user_password=password, owner_password=password)
        
        with open(output_path, 'wb') as f:
            writer.write(f)
        
        return output_path
    except Exception as e:
        logging.error(f"PDF parol qo'yishda xato: {e}")
        raise


def create_password_pdf(text: str, output_path: str, password: str, image_path: str = None):
    """Parolli PDF hujjat yaratish."""
    # Avval oddiy PDF yaratish
    temp_path = output_path + ".temp.pdf"
    create_pdf_document(text, temp_path, image_path=image_path)
    
    # Keyin parol qo'yish
    _add_password_to_pdf(temp_path, output_path, password)
    
    # Temp faylni o'chirish
    try:
        os.remove(temp_path)
    except Exception:
        pass
    
    return output_path


def create_multi_password_pdf(texts: list, output_path: str, password: str, image_paths: list = None):
    """Ko'p rasmli parolli PDF hujjat yaratish."""
    temp_path = output_path + ".temp.pdf"
    create_multi_image_pdf(texts, temp_path, image_paths=image_paths)
    _add_password_to_pdf(temp_path, output_path, password)
    
    try:
        os.remove(temp_path)
    except Exception:
        pass
    
    return output_path


# ==================== RASMNI AS-IS PDF/WORD GA ====================

def create_image_as_pdf(image_path: str, output_path: str):
    """Rasmni o'zgartirmasdan PDF ichiga joylashtirish."""
    from PIL import Image
    
    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)
    
    img = Image.open(image_path)
    img_w_px, img_h_px = img.size
    img.close()
    
    # A4: 210x297 mm
    page_w, page_h = 210, 297
    margin = 10
    usable_w = page_w - 2 * margin
    usable_h = page_h - 2 * margin
    
    # Rasmni sahifaga moslashtirish
    img_aspect = img_w_px / img_h_px
    page_aspect = usable_w / usable_h
    
    if img_aspect >= page_aspect:
        display_w = usable_w
        display_h = display_w / img_aspect
    else:
        display_h = usable_h
        display_w = display_h * img_aspect
    
    x = (page_w - display_w) / 2
    y = (page_h - display_h) / 2
    
    pdf.add_page()
    pdf.image(image_path, x=x, y=y, w=display_w, h=display_h)
    pdf.output(output_path)
    return output_path


def create_image_as_word(image_path: str, output_path: str):
    """Rasmni o'zgartirmasdan Word ichiga joylashtirish."""
    from PIL import Image
    
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    img = Image.open(image_path)
    img_w_px, img_h_px = img.size
    img.close()
    
    max_w_cm = 18.0  # 21 - 1.5 - 1.5
    max_h_cm = 27.7  # 29.7 - 1 - 1
    
    img_aspect = img_w_px / img_h_px
    page_aspect = max_w_cm / max_h_cm
    
    if img_aspect >= page_aspect:
        display_w = max_w_cm
        display_h = display_w / img_aspect
    else:
        display_h = max_h_cm
        display_w = display_h * img_aspect
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=Cm(display_w), height=Cm(display_h))
    
    doc.save(output_path)
    return output_path


def create_multi_image_as_pdf(image_paths: list, output_path: str):
    """Ko'p rasmni o'zgartirmasdan PDF ga joylashtirish."""
    from PIL import Image
    
    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)
    
    page_w, page_h = 210, 297
    margin = 10
    usable_w = page_w - 2 * margin
    usable_h = page_h - 2 * margin
    
    for img_path in image_paths:
        try:
            img = Image.open(img_path)
            img_w_px, img_h_px = img.size
            img.close()
            
            img_aspect = img_w_px / img_h_px
            page_aspect = usable_w / usable_h
            
            if img_aspect >= page_aspect:
                display_w = usable_w
                display_h = display_w / img_aspect
            else:
                display_h = usable_h
                display_w = display_h * img_aspect
            
            x = (page_w - display_w) / 2
            y = (page_h - display_h) / 2
            
            pdf.add_page()
            pdf.image(img_path, x=x, y=y, w=display_w, h=display_h)
        except Exception as e:
            logging.warning(f"Rasm qo'shilmadi: {e}")
    
    pdf.output(output_path)
    return output_path


def create_multi_image_as_word(image_paths: list, output_path: str):
    """Ko'p rasmni o'zgartirmasdan Word ga joylashtirish."""
    from PIL import Image
    
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    max_w_cm = 18.0
    max_h_cm = 27.7
    
    for idx, img_path in enumerate(image_paths):
        if idx > 0:
            doc.add_page_break()
        
        try:
            img = Image.open(img_path)
            img_w_px, img_h_px = img.size
            img.close()
            
            img_aspect = img_w_px / img_h_px
            page_aspect = max_w_cm / max_h_cm
            
            if img_aspect >= page_aspect:
                display_w = max_w_cm
                display_h = display_w / img_aspect
            else:
                display_h = max_h_cm
                display_w = display_h * img_aspect
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            run.add_picture(img_path, width=Cm(display_w), height=Cm(display_h))
        except Exception as e:
            logging.warning(f"Rasm qo'shilmadi: {e}")
    
    doc.save(output_path)
    return output_path

