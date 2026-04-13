import asyncio
import logging
import os
import time
import html as html_module
from aiohttp import web
from aiogram import BaseMiddleware, Bot, Dispatcher, types, F
from aiogram.filters import Command, CommandStart, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from aiogram.types import (
    InlineKeyboardMarkup, InlineKeyboardButton,
    CallbackQuery, FSInputFile, BotCommand
)
from dotenv import load_dotenv

from utils import (
    process_image_async,
    convert_latin_to_cyrillic,
    convert_cyrillic_to_latin,
    detect_script,
    translate_text,
    process_voice_async,
    calculate_quality_score,
)
from document_builder import (
    create_word_document, create_pdf_document,
    create_multi_image_word, create_multi_image_pdf,
    create_password_pdf, create_multi_password_pdf,
    create_image_as_pdf, create_image_as_word,
    create_multi_image_as_pdf, create_multi_image_as_word,
)
import database

# ==================== SOZLAMALAR ====================

load_dotenv()
BOT_TOKEN = os.getenv("BOT_TOKEN")
ADMIN_ID = int(os.getenv("ADMIN_ID", "0"))

if not BOT_TOKEN:
    raise ValueError("BOT_TOKEN is missing in .env")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s"
)
logger = logging.getLogger("PDFBot")

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()
database.init_db()

class FileSizeMiddleware(BaseMiddleware):
    async def __call__(self, handler, event: types.Message, data: dict):
        size = 0
        if getattr(event, 'document', None): size = event.document.file_size
        elif getattr(event, 'photo', None): size = event.photo[-1].file_size
        elif getattr(event, 'voice', None): size = event.voice.file_size
        elif getattr(event, 'audio', None): size = event.audio.file_size
        elif getattr(event, 'video', None): size = event.video.file_size
        
        if size and size > 20 * 1024 * 1024:
            await event.answer("⚠️ <b>Kechirasiz!</b> Fayl hajmi Telegram ruxsatidan (20MB) oshib ketdi.\nIltimos, hajmi kichikroq fayl yuboring.", parse_mode="HTML")
            return
        return await handler(event, data)

dp.message.middleware(FileSizeMiddleware())

TEMP_DIR = "temp_files"
os.makedirs(TEMP_DIR, exist_ok=True)

MAX_MULTI_IMAGES = 100
WELCOME_BANNER = "welcome_banner.png"


# ==================== HOLATLAR ====================

class DocState(StatesGroup):
    main_menu = State()
    # Bitta rasm
    waiting_for_image = State()
    waiting_for_format = State()
    waiting_for_alphabet = State()
    waiting_for_password = State()
    # Ko'p rasm
    multi_count = State()
    collecting_images = State()
    multi_format = State()
    multi_alphabet = State()
    multi_waiting_for_password = State()
    # Matn konvertor
    waiting_for_text = State()
    # Tarjima
    waiting_for_translate_image = State()
    translate_select_lang = State()
    translate_select_alphabet = State()
    # Ovozli xabar
    voice_select_alphabet = State()
    voice_select_format = State()
    # Rasmni PDF/Word (asl rasm sifatida)
    waiting_for_raw_image = State()
    raw_select_format = State()
    # Ko'p rasmni PDF/Word (asl rasm)
    multi_raw_count = State()
    collecting_raw_images = State()
    multi_raw_format = State()
    # PDF <-> Word konvertor
    waiting_for_pdf_file = State()
    waiting_for_word_file = State()
    convert_select_alphabet = State()


# ==================== KLAVIATURALAR ====================

def get_main_menu_keyboard(is_admin=False):
    keyboard = [
        [InlineKeyboardButton(text="📷 Rasmni hujjatga", callback_data="mode_image")],
        [InlineKeyboardButton(text="📄 PDF → Word", callback_data="mode_pdf_to_word"),
         InlineKeyboardButton(text="📝 Word → PDF", callback_data="mode_word_to_pdf")],
        [InlineKeyboardButton(text="🌐 Rasmni tarjima + hujjat", callback_data="mode_translate")],
        [InlineKeyboardButton(text="🔤 Lotin ↔ Kirill", callback_data="mode_convert"),
         InlineKeyboardButton(text="🎤 Ovoz", callback_data="mode_voice")],
        [InlineKeyboardButton(text="⚙️ Hisobim & Yordam", callback_data="account_help_menu")]
    ]
    return InlineKeyboardMarkup(inline_keyboard=keyboard)

def get_account_help_keyboard(is_admin=False):
    keyboard = [
        [InlineKeyboardButton(text="📊 Hisobim", callback_data="myaccount"),
         InlineKeyboardButton(text="❓ Yordam", callback_data="help_menu")]
    ]
    if is_admin:
        # Reyting faqat adminga ko'rinadi
        keyboard.append([InlineKeyboardButton(text="🏆 Reyting", callback_data="leaderboard")])
    keyboard.append([InlineKeyboardButton(text="◀️ Orqaga", callback_data="main_menu")])
    return InlineKeyboardMarkup(inline_keyboard=keyboard)

def get_raw_format_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="📕 PDF", callback_data="rawfmt_PDF"),
            InlineKeyboardButton(text="📘 Word", callback_data="rawfmt_Word")
        ],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="main_menu")]
    ])

def get_format_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="📕 PDF", callback_data="fmt_PDF"),
            InlineKeyboardButton(text="📘 Word", callback_data="fmt_Word")
        ],
        [
            InlineKeyboardButton(text="🔒 Parolli PDF", callback_data="fmt_PassPDF"),
            InlineKeyboardButton(text="📋 Faqat matn", callback_data="fmt_Text")
        ],
        [
            InlineKeyboardButton(text="🖼 Rasm sifatida PDF", callback_data="fmt_RawPDF"),
            InlineKeyboardButton(text="🖼 Rasm sifatida Word", callback_data="fmt_RawWord")
        ]
    ])

def get_alphabet_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="🔤 Lotin", callback_data="abc_Lotin"),
            InlineKeyboardButton(text="🔤 Кирилл", callback_data="abc_Kirill")
        ]
    ])

def get_translate_lang_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🇺🇿→🇷🇺 O'zbek → Rus", callback_data="tr_uz_ru")],
        [InlineKeyboardButton(text="🇷🇺→🇺🇿 Rus → O'zbek", callback_data="tr_ru_uz")],
        [InlineKeyboardButton(text="🇺🇿→🇬🇧 O'zbek → Ingliz", callback_data="tr_uz_en")],
        [InlineKeyboardButton(text="🇬🇧→🇺🇿 Ingliz → O'zbek", callback_data="tr_en_uz")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="main_menu")],
    ])

def get_feedback_keyboard(doc_hash: str):
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="👍 Zo'r", callback_data=f"fb_good_{doc_hash}"),
            InlineKeyboardButton(text="👎 Yaxshilansin", callback_data=f"fb_bad_{doc_hash}"),
        ],
        [InlineKeyboardButton(text="🏠 Bosh menyu", callback_data="main_menu")],
    ])

def get_after_doc_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏠 Bosh menyu", callback_data="main_menu")],
    ])

def _get_streak_emoji(streak: int) -> str:
    if streak >= 30: return "💎"
    if streak >= 14: return "🔥🔥"
    if streak >= 7: return "🔥"
    if streak >= 3: return "⚡"
    return "✨"



# ==================== /start ====================

@dp.message(CommandStart(), StateFilter("*"))
async def start_cmd(message: types.Message, state: FSMContext):
    await state.clear()
    is_new = database.register_user(
        user_id=message.from_user.id,
        first_name=message.from_user.first_name,
        username=message.from_user.username,
    )
    
    # Streak yangilash
    streak_info = database.update_streak(message.from_user.id)
    
    if is_new:
        await _send_welcome_guide(message, state)
    else:
        if streak_info.get("new_bonus"):
            bonus = streak_info["bonus"]
            streak = streak_info["streak"]
            await message.answer(
                f"🔥 <b>Streak bonus!</b>\n\n"
                f"{_get_streak_emoji(streak)} {streak} kun ketma-ket!\n"
                f"🎁 +{bonus} ta bepul imkoniyat berildi!",
                parse_mode="HTML"
            )
        await _send_main_menu(message, state)


async def _send_welcome_guide(message: types.Message, state: FSMContext):
    """Yangi foydalanuvchi uchun chiroyli kirish ko'rsatmasi."""
    name = message.from_user.first_name or "Foydalanuvchi"
    
    guide = (
        f"🎉 <b>Xush kelibsiz, {html_module.escape(name)}!</b>\n\n"
        f"Men — <b>AI Hujjat Bot</b> 🤖\n"
        f"Rasmingizni yuborasiz, men uni chiroyli\n"
        f"PDF yoki Word hujjatiga aylantirib beraman!\n\n"
        f"━━━━━━━━━━━━━━━━━━━━\n"
        f"🔰 <b>3 qadam bilan ishlaydi:</b>\n\n"
        f"1️⃣ Pastdagi tugmani bosing\n"
        f"     ↳ <i>Masalan: \"📸 Rasmni hujjatga aylantir\"</i>\n\n"
        f"2️⃣ Rasmni menga yuboring\n"
        f"     ↳ <i>Daftar, kitob, hujjat rasmi</i>\n\n"
        f"3️⃣ PDF yoki Word tanlang — tayyor! ✅\n\n"
        f"━━━━━━━━━━━━━━━━━━━━\n"
        f"🎁 Sizga <b>2 ta bepul</b> foydalanish berildi!\n"
        f"♻️ Har kuni yana 2 tadan yangilanadi\n\n"
        f"💡 <i>Batafsil: pastda \"❓ Qanday ishlaydi?\" bosing</i>"
    )
    
    # Banner rasm bilan yuborish
    try:
        if os.path.exists(WELCOME_BANNER):
            banner = FSInputFile(WELCOME_BANNER)
            await message.answer_photo(
                photo=banner,
                caption=guide,
                parse_mode="HTML"
            )
        else:
            await message.answer(guide, parse_mode="HTML")
    except Exception:
        await message.answer(guide, parse_mode="HTML")
    
    # Keyin asosiy menyu
    await asyncio.sleep(0.5)
    await _send_main_menu(message, state)


async def _send_main_menu(message: types.Message, state: FSMContext):
    uid = message.from_user.id if message.from_user else message.chat.id
    name = message.from_user.first_name if message.from_user else "Foydalanuvchi"
    is_admin = (ADMIN_ID and uid == ADMIN_ID)
    balance = "♾️" if is_admin else database.get_user_balance(uid)
    streak = database.get_streak(uid)
    streak_emoji = _get_streak_emoji(streak)
    
    welcome = (
        f"👋 Salom, <b>{html_module.escape(name)}</b>!\n"
        f"🎯 Bepul so'rovlar: <b>{balance}/{database.DEFAULT_DAILY_LIMIT}</b>\n"
        f"⏰ Yangilanadi: ertaga 00:00 da\n"
        f"👇 Nima qilmoqchisiz?"
    )
    await message.answer(welcome, reply_markup=get_main_menu_keyboard(is_admin), parse_mode="HTML")
    await state.set_state(DocState.main_menu)


# ==================== BOSH MENYU CALLBACK ====================

@dp.callback_query(F.data == "main_menu")
async def main_menu_callback(callback: CallbackQuery, state: FSMContext):
    await state.clear()
    uid = callback.from_user.id
    name = callback.from_user.first_name or "Foydalanuvchi"
    is_admin = (ADMIN_ID and uid == ADMIN_ID)
    balance = "♾️" if is_admin else database.get_user_balance(uid)
    streak = database.get_streak(uid)
    streak_emoji = _get_streak_emoji(streak)
    
    welcome = (
        f"👋 Salom, <b>{html_module.escape(name)}</b>!\n"
        f"🎯 Bepul so'rovlar: <b>{balance}/{database.DEFAULT_DAILY_LIMIT}</b>\n"
        f"⏰ Yangilanadi: ertaga 00:00 da\n"
        f"👇 Nima qilmoqchisiz?"
    )
    try:
        await callback.message.edit_text(welcome, reply_markup=get_main_menu_keyboard(is_admin), parse_mode="HTML")
    except Exception:
        await callback.message.answer(welcome, reply_markup=get_main_menu_keyboard(is_admin), parse_mode="HTML")
    await state.set_state(DocState.main_menu)
    await callback.answer()


# ==================== REJIM TANLASH ====================

@dp.callback_query(F.data == "mode_raw_image")
async def mode_raw_image(callback: CallbackQuery, state: FSMContext):
    await callback.message.edit_text(
        "🖼 <b>Rasmni PDF/Word</b>\n\n"
        "Rasmni yuboring — matn ajratilmaydi, faqat rasmning o'zini PDF/Word hujjatiga aylantiraman.\n"
        "💡 <i>Sifatli rasm uchun uni fayl sifatida yuborishingiz mumkin.</i>",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.waiting_for_raw_image)
    await callback.answer()

@dp.callback_query(F.data == "mode_multi_raw")
async def mode_multi_raw(callback: CallbackQuery, state: FSMContext):
    await state.update_data(multi_raw_images=[])
    await callback.message.edit_text(
        f"🖼🖼 <b>Ko'p Asl Rasm → PDF/Word</b>\n\n"
        f"Nechta varaq (rasm) yuborasiz?\n"
        f"Sonni yozing (1 dan {MAX_MULTI_IMAGES} gacha):\n\n"
        f"💡 <i>Masalan: 3</i>",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.multi_raw_count)
    await callback.answer()


@dp.callback_query(F.data == "mode_image")
async def mode_image(callback: CallbackQuery, state: FSMContext):
    await state.update_data(multi_images=[])
    await callback.message.edit_text(
        f"📸 <b>Rasmni hujjatga aylantirish</b>\n"
        f"━━━━━━━━━━━━━━━━━━━━\n\n"
        f"Nechta rasm yuborasiz?\n\n"
        f"👇 Tugmani bosing yoki boshqa son\n"
        f"uchun raqam <b>yozing</b> (1-{MAX_MULTI_IMAGES}):\n\n"
        f"💡 <i>Masalan: 7 deb yozing va yuboring</i>",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="1️⃣ 1 ta rasm", callback_data="imgcount_1"),
             InlineKeyboardButton(text="2️⃣ 2 ta", callback_data="imgcount_2"),
             InlineKeyboardButton(text="3️⃣ 3 ta", callback_data="imgcount_3")],
            [InlineKeyboardButton(text="5️⃣ 5 ta", callback_data="imgcount_5"),
             InlineKeyboardButton(text="🔟 10 ta", callback_data="imgcount_10")],
            [InlineKeyboardButton(text="✍️ Boshqa son yozish", callback_data="imgcount_custom")],
            [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.multi_count)
    await callback.answer()


@dp.callback_query(DocState.multi_count, F.data.startswith("imgcount_"))
async def imgcount_quick_callback(callback: CallbackQuery, state: FSMContext):
    """Tezkor tugma orqali rasm sonini tanlash."""
    data_val = callback.data.split("_")[1]
    if data_val == "custom":
        await callback.message.edit_reply_markup(reply_markup=None)
        await callback.message.answer(f"✍️ <b>Raqam yozing (1-{MAX_MULTI_IMAGES})</b>\n\nMasalan: <code>7</code> deb yozib yuboring", parse_mode="HTML")
        await callback.answer()
        return
        
    count = int(data_val)
    await _setup_image_collection(callback.message, state, count, edit=True)
    await callback.answer()


@dp.message(DocState.multi_count, F.text)
async def multi_count_handler(message: types.Message, state: FSMContext):
    """Foydalanuvchi nechta rasm yuborishini kiritadi."""
    if message.text.startswith('/'):
        return
    
    text = message.text.strip()
    if not text.isdigit():
        await message.answer(
            f"⚠️ Iltimos, faqat son kiriting (1-{MAX_MULTI_IMAGES}).",
            parse_mode="HTML"
        )
        return
    
    count = int(text)
    if count < 1 or count > MAX_MULTI_IMAGES:
        await message.answer(
            f"⚠️ Son 1 dan {MAX_MULTI_IMAGES} gacha bo'lishi kerak.",
            parse_mode="HTML"
        )
        return
    
    await _setup_image_collection(message, state, count, edit=False)


async def _setup_image_collection(message_or_msg, state: FSMContext, count: int, edit: bool = False):
    """Son kiritilgandan keyin rasm kutish holatini sozlash."""
    await state.update_data(multi_count=count, multi_images=[])
    
    if count == 1:
        text = (
            "📸 <b>1 ta rasm</b>\n"
            "━━━━━━━━━━━━━━━━━━━━\n\n"
            "Rasmni menga yuboring\n"
            "<i>(daftar, kitob, hujjat rasmi...)</i>\n\n"
            "👇 <b>Rasmni hozir yuboring:</b>"
        )
        next_state = DocState.waiting_for_image
    else:
        text = (
            f"✅ <b>{count} ta rasm</b> kutilmoqda.\n\n"
            f"📸 Rasmlarni birma-bir yuboring.\n"
            f"Yuborilgan: <b>0</b> / {count}"
        )
        next_state = DocState.collecting_images
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="main_menu")],
    ])
    
    if edit:
        try:
            await message_or_msg.edit_text(text, reply_markup=kb, parse_mode="HTML")
        except Exception:
            await message_or_msg.answer(text, reply_markup=kb, parse_mode="HTML")
    else:
        await message_or_msg.answer(text, reply_markup=kb, parse_mode="HTML")
    
    await state.set_state(next_state)


@dp.callback_query(F.data == "mode_translate")
async def mode_translate(callback: CallbackQuery, state: FSMContext):
    await callback.message.edit_text(
        "🌐 <b>Rasmni tarjima qilib hujjat</b>\n"
        "━━━━━━━━━━━━━━━━━━━━\n\n"
        "Rasmni yuboring — men undagi matnni\n"
        "boshqa tilga tarjima qilib, PDF beraman!\n\n"
        "🇺🇿→🇷🇺 O'zbek → Rus\n"
        "🇷🇺→🇺🇿 Rus → O'zbek\n"
        "🇺🇿→🇬🇧 O'zbek → Ingliz\n"
        "🇬🇧→🇺🇿 Ingliz → O'zbek\n\n"
        "👇 <b>Rasmni hozir yuboring:</b>",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.waiting_for_translate_image)
    await callback.answer()


@dp.message(DocState.waiting_for_translate_image, F.photo)
async def translate_image_photo(message: types.Message, state: FSMContext):
    """Tarjima uchun rasm qabul qilish."""
    is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
    if not is_admin:
        has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
        if not has_limit:
            await message.answer(
                "⚡ <b>Bugungi limitingiz tugadi!</b>\n"
                "⏰ Ertaga soat 00:00 da yangilanadi\n"
                "💎 Premium olish → /premium",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            return
    
    photo_file = await bot.get_file(message.photo[-1].file_id)
    image_path = os.path.join(TEMP_DIR, f"tr_{message.from_user.id}_{int(time.time())}.jpg")
    await bot.download_file(photo_file.file_path, image_path)
    
    await state.update_data(image_path=image_path)
    await message.answer(
        "✅ Rasm qabul qilindi!\n\n"
        "🌐 Qaysi tilga tarjima qilish kerak?",
        reply_markup=get_translate_lang_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(DocState.translate_select_lang)


@dp.message(DocState.waiting_for_translate_image)
async def translate_not_image(message: types.Message):
    if message.text and message.text.startswith('/'):
        return
    await message.answer("📸 Iltimos, rasmni yuboring!", parse_mode="HTML")


@dp.callback_query(DocState.translate_select_lang, F.data.startswith("tr_"))
async def translate_lang_callback(callback: CallbackQuery, state: FSMContext):
    await state.set_state('processing')
    lang = callback.data[3:]  # uz_ru, ru_uz, uz_en, en_uz
    await state.update_data(translate_lang=lang)
    
    lang_labels = {
        "uz_ru": "🇺🇿→🇷🇺 O'zbek→Rus",
        "ru_uz": "🇷🇺→🇺🇿 Rus→O'zbek",
        "uz_en": "🇺🇿→🇬🇧 O'zbek→Ingliz",
        "en_uz": "🇬🇧→🇺🇿 Ingliz→O'zbek",
    }
    label = lang_labels.get(lang, lang)
    
    await callback.message.edit_text(
        f"🌐 Tarjima: <b>{label}</b> ✅\n\n"
        f"Qaysi alifboda chiqarish kerak?",
        reply_markup=get_alphabet_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(DocState.translate_select_alphabet)
    await callback.answer()


@dp.callback_query(DocState.translate_select_alphabet, F.data.startswith("abc_"))
async def translate_alphabet_callback(callback: CallbackQuery, state: FSMContext):
    await state.set_state('processing')
    selected_alphabet = callback.data.split("_")[1]
    data = await state.get_data()
    image_path = data.get("image_path")
    translate_lang = data.get("translate_lang")
    
    if not image_path or not os.path.exists(image_path):
        await callback.message.edit_text("❌ Rasm topilmadi. /start bosing.")
        await state.clear()
        await callback.answer()
        return
    
    status_msg = await callback.message.edit_text(
        "🌐 <b>Tarjima + Hujjat</b>\n\n"
        "⏳ <i>Rasm tahlil qilinmoqda...</i>\n"
        "▓░░░░░░░░░ 10%",
        parse_mode="HTML"
    )
    await callback.answer()
    
    try:
        # AI tahlil
        try:
            await callback.message.edit_text(
                "🌐 <b>Tarjima + Hujjat</b>\n\n"
                "🔍 <i>AI matnni aniqlayapti...</i>\n"
                "▓▓▓░░░░░░░ 30%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        extracted_text, cropped_path = await process_image_async(image_path, selected_alphabet)
        
        if not extracted_text or len(extracted_text.strip()) < 3:
            await callback.message.edit_text(
                "⚠️ <b>Rasmda matn topilmadi</b>",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            _cleanup_file(image_path)
            await state.clear()
            await state.set_state(DocState.main_menu)
            return
        
        # Tarjima
        try:
            await callback.message.edit_text(
                "🌐 <b>Tarjima + Hujjat</b>\n\n"
                "🔄 <i>Tarjima qilinmoqda...</i>\n"
                "▓▓▓▓▓▓░░░░ 60%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        translated_text = await translate_text(extracted_text, translate_lang)
        
        # Hujjat yaratish
        try:
            await callback.message.edit_text(
                "🌐 <b>Tarjima + Hujjat</b>\n\n"
                "📝 <i>PDF yaratilmoqda...</i>\n"
                "▓▓▓▓▓▓▓▓░░ 80%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        base_name = f"tr_{callback.from_user.id}_{int(time.time())}"
        output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
        await asyncio.to_thread(create_pdf_document, translated_text, output_path, cropped_path)
        doc_file = FSInputFile(output_path, filename="Tarjima.pdf")
        
        try:
            await callback.message.edit_text(
                "🌐 <b>Tarjima + Hujjat</b>\n\n"
                "✅ <i>Tayyor!</i>\n"
                "▓▓▓▓▓▓▓▓▓▓ 100%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        lang_labels = {"uz_ru": "O'zbek→Rus", "ru_uz": "Rus→O'zbek", "uz_en": "O'zbek→Ingliz", "en_uz": "Ingliz→O'zbek"}
        word_count = len(translated_text.split())
        
        doc_hash = database.generate_doc_hash(translated_text, callback.from_user.id)
        database.save_doc_hash(doc_hash, callback.from_user.id, f"Tarjima ({lang_labels.get(translate_lang, '')})", word_count)
        
        await callback.message.answer_document(
            document=doc_file,
            caption=(
                f"✅ <b>Tarjima tayyor!</b>\n\n"
                f"🌐 Yo'nalish: <b>{lang_labels.get(translate_lang, '')}</b>\n"
                f"📄 Format: <b>PDF</b>\n"
                f"📊 ~{word_count} so'z\n"
                f"🔐 ID: <code>{doc_hash}</code>"
            ),
            parse_mode="HTML"
        )
        await callback.message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
        
        _cleanup_file(image_path)
        _cleanup_file(output_path)
        if 'cropped_path' in locals() and cropped_path:
            _cleanup_file(cropped_path)
        
    except Exception as e:
        logger.error(f"Tarjima xatolik: {e}", exc_info=True)
        error_msg = "❌ <b>Xatolik yuz berdi</b>\n\n🔄 Qayta urinib ko'ring."
        if ADMIN_ID and callback.from_user.id == ADMIN_ID:
            error_msg += f"\n\n🔧 <code>{html_module.escape(str(e)[:500])}</code>"
        await callback.message.answer(error_msg, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")
        _cleanup_file(image_path)
        if 'cropped_path' in locals() and cropped_path:
            _cleanup_file(cropped_path)
        if 'output_path' in locals() and output_path:
            _cleanup_file(output_path)
    
    finally:
        await state.clear()
        await state.set_state(DocState.main_menu)


@dp.callback_query(F.data == "mode_convert")
async def mode_convert(callback: CallbackQuery, state: FSMContext):
    await callback.message.edit_text(
        "🔄 <b>Lotin ↔ Кирилл o'girish</b>\n"
        "━━━━━━━━━━━━━━━━━━━━\n\n"
        "Matn yozing — avtomatik aniqlayman:\n"
        "• Lotin yozsangiz → Кирилл ga o'giraman\n"
        "• Кирилл yozsangiz → Lotin ga o'giraman\n\n"
        "👇 <b>Matnni hozir yozing va yuboring:</b>",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.waiting_for_text)
    await callback.answer()


# ==================== MATN KONVERTOR ====================

@dp.message(DocState.waiting_for_text, F.text)
async def convert_text_handler(message: types.Message, state: FSMContext):
    if message.text.startswith('/'):
        return
    
    text = message.text
    script = detect_script(text)
    
    if script == "kirill":
        converted = convert_cyrillic_to_latin(text)
        direction = "Кирилл → Lotin"
    else:
        converted = convert_latin_to_cyrillic(text)
        direction = "Lotin → Кирилл"
    
    # HTML injection himoyasi
    safe_converted = html_module.escape(converted)
    
    result = (
        f"🔄 <b>{direction}</b>\n\n"
        f"━━━━━━━━━━━━━━━━━━━━\n"
        f"{safe_converted}\n"
        f"━━━━━━━━━━━━━━━━━━━━\n\n"
        f"✏️ Yana matn yuboring yoki ⬇️"
    )
    
    await message.answer(
        result, 
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Bosh menyu", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )

# ==================== RASMNI AS-IS PDF/WORD GA AYLANTRISH ====================

@dp.message(DocState.waiting_for_raw_image, F.photo)
async def raw_image_photo(message: types.Message, state: FSMContext):
    """Rasm qabul qilish."""
    is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
    
    if not is_admin:
        has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
        if not has_limit:
            await message.answer(
                "⚠️ <b>Bugungi limitingiz tugadi!</b>\n💡 Ertaga yana 2 ta tekin beriladi.",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            return
    else:
        remaining = "♾️"
        
    photo_file = await bot.get_file(message.photo[-1].file_id)
    image_path = os.path.join(TEMP_DIR, f"raw_{message.from_user.id}_{message.message_id}_{int(time.time())}.jpg")
    await bot.download_file(photo_file.file_path, image_path)
    
    await state.update_data(image_path=image_path)
    await message.answer(
        f"✅ Rasm qabul qilindi! (Qoldiq: <b>{remaining}</b>)\n\n"
        f"Qaysi formatda hujjat olmoqchisiz?",
        reply_markup=get_raw_format_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(DocState.raw_select_format)

@dp.message(DocState.waiting_for_raw_image, F.document)
async def raw_image_document(message: types.Message, state: FSMContext):
    doc = message.document
    if not doc.mime_type or not doc.mime_type.startswith("image/"):
        await message.answer(
            "⚠️ Bu rasm fayli emas. Iltimos, rasm yuboring.",
            parse_mode="HTML"
        )
        return
    
    is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
    
    if not is_admin:
        has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
        if not has_limit:
            await message.answer(
                "⚠️ <b>Bugungi limitingiz tugadi!</b>\n💡 Ertaga yana 2 ta tekin beriladi.",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            return
    else:
        remaining = "♾️"
        
    file = await bot.get_file(doc.file_id)
    ext = os.path.splitext(doc.file_name or "img.jpg")[1] or ".jpg"
    image_path = os.path.join(TEMP_DIR, f"raw_{message.from_user.id}_{message.message_id}_{int(time.time())}{ext}")
    await bot.download_file(file.file_path, image_path)
    
    await state.update_data(image_path=image_path)
    await message.answer(
        f"✅ Rasm qabul qilindi! (Qoldiq: <b>{remaining}</b>)\n\n"
        f"Qaysi formatda hujjat olmoqchisiz?",
        reply_markup=get_raw_format_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(DocState.raw_select_format)

@dp.message(DocState.waiting_for_raw_image)
async def raw_not_image(message: types.Message):
    if message.text and message.text.startswith('/'):
        return
    await message.answer("📸 Iltimos, rasmni yuboring.", parse_mode="HTML")

@dp.callback_query(DocState.raw_select_format, F.data.startswith("rawfmt_"))
async def raw_format_callback(callback: CallbackQuery, state: FSMContext):
    await state.set_state('processing')
    selected_format = callback.data.split("_", 1)[1]
    data = await state.get_data()
    image_path = data.get("image_path")
    user_id = callback.from_user.id
    
    if not image_path or not os.path.exists(image_path):
        await callback.message.edit_text("❌ Rasm topilmadi. /start bosing.")
        await state.clear()
        await callback.answer()
        return
        
    fmt_emoji = "📕" if selected_format == "PDF" else "📘"
    
    status_msg = await callback.message.edit_text(
        f"{fmt_emoji} {selected_format} (Asl rasm)\n\n"
        f"⏳ <i>Hujjat yaratilmoqda...</i>\n"
        f"▓▓▓▓▓░░░░░ 50%",
        parse_mode="HTML"
    )
    await callback.answer()
    
    try:
        base_name = f"doc_{user_id}_{int(time.time())}"
        
        if selected_format == "PDF":
            output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
            await asyncio.to_thread(create_image_as_pdf, image_path, output_path)
            doc_file = FSInputFile(output_path, filename="Hujjat.pdf")
        else:
            output_path = os.path.join(TEMP_DIR, f"{base_name}.docx")
            await asyncio.to_thread(create_image_as_word, image_path, output_path)
            doc_file = FSInputFile(output_path, filename="Hujjat.docx")
            
        try:
            await status_msg.edit_text(
                f"{fmt_emoji} {selected_format} (Asl rasm)\n\n"
                f"✅ <i>Tayyor! Yuborilmoqda...</i>\n"
                f"▓▓▓▓▓▓▓▓▓▓ 100%",
                parse_mode="HTML"
            )
        except Exception:
            pass
            
        doc_hash = database.generate_doc_hash(f"raw_img_{time.time()}", user_id)
        database.save_doc_hash(doc_hash, user_id, f"{selected_format} (Rasm)", 0)
        
        caption = (
            f"✅ <b>Hujjat tayyor!</b>\n\n"
            f"📄 Format: <b>{selected_format}</b> (Rasm)\n"
            f"🔐 ID: <code>{doc_hash}</code>"
        )
        await callback.message.answer_document(
            document=doc_file,
            caption=caption,
            parse_mode="HTML"
        )
        await callback.message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
        
        _cleanup_file(image_path)
        _cleanup_file(output_path)
        
    except Exception as e:
        logger.error(f"Xatolik: {e}", exc_info=True)
        error_msg = "❌ <b>Xatolik yuz berdi</b>\n\n🔄 Iltimos, qayta urinib ko'ring."
        await callback.message.answer(error_msg, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")
        _cleanup_file(image_path)
        if 'output_path' in locals() and output_path:
            _cleanup_file(output_path)
    finally:
        await state.clear()
        await state.set_state(DocState.main_menu)


# ==================== BITTA RASM — QABUL QILISH ====================

@dp.message(DocState.waiting_for_image, F.photo)
async def single_image_photo(message: types.Message, state: FSMContext):
    """Rasm (photo) qabul qilish."""
    is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
    
    if not is_admin:
        has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
        if not has_limit:
            await message.answer(
                "⚠️ <b>Bugungi limitingiz tugadi!</b>\n💡 Ertaga yana 2 ta tekin beriladi.",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            return
    else:
        remaining = "♾️"
    
    photo_file = await bot.get_file(message.photo[-1].file_id)
    image_path = os.path.join(TEMP_DIR, f"{message.from_user.id}_{message.message_id}_{int(time.time())}.jpg")
    await bot.download_file(photo_file.file_path, image_path)
    
    await state.update_data(image_path=image_path)
    await message.answer(
        f"✅ Rasm qabul qilindi! (Qoldiq: <b>{remaining}</b>)\n\n"
        f"Qaysi formatda hujjat olmoqchisiz?",
        reply_markup=get_format_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(DocState.waiting_for_format)


@dp.message(DocState.waiting_for_image, F.document)
async def single_image_document(message: types.Message, state: FSMContext):
    """Fayl (document) sifatida yuborilgan rasmni qabul qilish."""
    doc = message.document
    if not doc.mime_type or not doc.mime_type.startswith("image/"):
        await message.answer(
            "⚠️ Bu rasm fayli emas. Iltimos, rasm yuboring.",
            parse_mode="HTML"
        )
        return
    
    is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
    
    if not is_admin:
        has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
        if not has_limit:
            await message.answer(
                "⚠️ <b>Bugungi limitingiz tugadi!</b>\n💡 Ertaga yana 2 ta tekin beriladi.",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            return
    else:
        remaining = "♾️"
    
    file = await bot.get_file(doc.file_id)
    ext = os.path.splitext(doc.file_name or "img.jpg")[1] or ".jpg"
    image_path = os.path.join(TEMP_DIR, f"{message.from_user.id}_{message.message_id}_{int(time.time())}{ext}")
    await bot.download_file(file.file_path, image_path)
    
    await state.update_data(image_path=image_path)
    await message.answer(
        f"✅ Rasm qabul qilindi! (Qoldiq: <b>{remaining}</b>)\n\n"
        f"Qaysi formatda hujjat olmoqchisiz?",
        reply_markup=get_format_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(DocState.waiting_for_format)


@dp.message(DocState.waiting_for_image)
async def single_not_image(message: types.Message):
    if message.text and message.text.startswith('/'):
        return
    await message.answer("📸 Iltimos, rasmni yuboring.", parse_mode="HTML")


# ==================== BITTA RASM — FORMAT VA ALIFBO ====================

@dp.callback_query(DocState.waiting_for_format, F.data.startswith("fmt_"))
async def single_format_callback(callback: CallbackQuery, state: FSMContext):
    selected_format = callback.data.split("_", 1)[1]
    
    # Rasm sifatida (matn ajratilmaydi) — alohida oqim
    if selected_format in ("RawPDF", "RawWord"):
        await state.set_state('processing')
        data = await state.get_data()
        image_path = data.get("image_path")
        user_id = callback.from_user.id
        
        if not image_path or not os.path.exists(image_path):
            await callback.message.edit_text("❌ Rasm topilmadi. /start bosing.")
            await state.clear()
            await callback.answer()
            return
        
        actual_format = "PDF" if selected_format == "RawPDF" else "Word"
        fmt_emoji = "📕" if actual_format == "PDF" else "📘"
        
        status_msg = await callback.message.edit_text(
            f"🖼 {actual_format} (Rasm sifatida)\n\n"
            f"⏳ <i>Hujjat yaratilmoqda...</i>\n"
            f"▓▓▓▓▓░░░░░ 50%",
            parse_mode="HTML"
        )
        await callback.answer()
        
        try:
            base_name = f"doc_{user_id}_{int(time.time())}"
            if actual_format == "PDF":
                output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
                await asyncio.to_thread(create_image_as_pdf, image_path, output_path)
                doc_file = FSInputFile(output_path, filename="Hujjat.pdf")
            else:
                output_path = os.path.join(TEMP_DIR, f"{base_name}.docx")
                await asyncio.to_thread(create_image_as_word, image_path, output_path)
                doc_file = FSInputFile(output_path, filename="Hujjat.docx")
            
            try:
                await status_msg.edit_text(
                    f"🖼 {actual_format} (Rasm sifatida)\n\n"
                    f"✅ <i>Tayyor!</i>\n"
                    f"▓▓▓▓▓▓▓▓▓▓ 100%",
                    parse_mode="HTML"
                )
            except Exception:
                pass
            
            doc_hash = database.generate_doc_hash(f"raw_img_{time.time()}", user_id)
            database.save_doc_hash(doc_hash, user_id, f"{actual_format} (Rasm)", 0)
            
            await callback.message.answer_document(
                document=doc_file,
                caption=(
                    f"✅ <b>Hujjat tayyor!</b>\n\n"
                    f"📄 Format: <b>{actual_format}</b> (Rasm sifatida)\n"
                    f"🔐 ID: <code>{doc_hash}</code>"
                ),
                parse_mode="HTML"
            )
            await callback.message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
            
            _cleanup_file(image_path)
            _cleanup_file(output_path)
            
        except Exception as e:
            logger.error(f"Raw rasm xatolik: {e}", exc_info=True)
            await callback.message.answer(
                "❌ <b>Xatolik yuz berdi</b>\n\n🔄 Qayta urinib ko'ring.",
                reply_markup=get_after_doc_keyboard(), parse_mode="HTML"
            )
            _cleanup_file(image_path)
            if 'output_path' in locals() and output_path:
                _cleanup_file(output_path)
        finally:
            await state.clear()
            await state.set_state(DocState.main_menu)
        return
    
    # Oddiy matn ajratish oqimi
    await state.update_data(format=selected_format)
    
    fmt_map = {"PDF": "📕", "Word": "📘", "PassPDF": "🔒", "Text": "📋"}
    fmt_emoji = fmt_map.get(selected_format, "📄")
    fmt_label = "Parolli PDF" if selected_format == "PassPDF" else ("Faqat matn" if selected_format == "Text" else selected_format)
    
    await callback.message.edit_text(
        f"{fmt_emoji} Format: <b>{fmt_label}</b> ✅\n\n"
        f"Qaysi alifboda bo'lsin?",
        reply_markup=get_alphabet_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(DocState.waiting_for_alphabet)
    await callback.answer()


@dp.callback_query(DocState.waiting_for_alphabet, F.data.startswith("abc_"))
async def single_alphabet_callback(callback: CallbackQuery, state: FSMContext):
    await state.set_state('processing')
    selected_alphabet = callback.data.split("_")[1]
    await state.update_data(alphabet=selected_alphabet)
    data = await state.get_data()
    image_path = data.get("image_path")
    selected_format = data.get("format")
    
    if not image_path or not os.path.exists(image_path):
        await callback.message.edit_text("❌ Rasm topilmadi. /start bosing.")
        await state.clear()
        await callback.answer()
        return
    
    # PassPDF bo'lsa — avval parolni so'rash
    if selected_format == "PassPDF":
        await callback.message.edit_text(
            "🔒 <b>Parolli PDF</b>\n\n"
            "PDF uchun parolni yozing:\n"
            "💡 <i>Kamida 4 belgi</i>",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="main_menu")],
            ]),
            parse_mode="HTML"
        )
        await state.set_state(DocState.waiting_for_password)
        await callback.answer()
        return
    
    await callback.answer()
    await _process_single_image(callback.message, state, callback.from_user.id)


async def _process_single_image(message, state: FSMContext, user_id: int, password: str = None):
    """Bitta rasmni qayta ishlash — barcha formatlar uchun."""
    data = await state.get_data()
    image_path = data.get("image_path")
    selected_format = data.get("format")
    selected_alphabet = data.get("alphabet")
    
    if selected_format == "PassPDF":
        fmt_emoji = "🔒"
        fmt_label = "Parolli PDF"
    elif selected_format == "Text":
        fmt_emoji = "📋"
        fmt_label = "Matn"
    else:
        fmt_emoji = "📕" if selected_format == "PDF" else "📘"
        fmt_label = selected_format
    
    # 1-bosqich
    status_msg = await message.answer(
        f"{fmt_emoji} {fmt_label} | 🔤 {selected_alphabet}\n\n"
        f"⏳ <i>Rasm tahlil qilinmoqda...</i>\n"
        f"▓░░░░░░░░░ 10%",
        parse_mode="HTML"
    )
    
    try:
        # 2-bosqich: AI tahlil
        try:
            await status_msg.edit_text(
                f"{fmt_emoji} {fmt_label} | 🔤 {selected_alphabet}\n\n"
                f"🔍 <i>AI matnni aniqlayapti...</i>\n"
                f"▓▓▓▓░░░░░░ 40%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        extracted_text, cropped_path = await process_image_async(image_path, selected_alphabet)
        
        if not extracted_text or len(extracted_text.strip()) < 3:
            await status_msg.edit_text(
                "⚠️ <b>Rasmda matn topilmadi</b>\n\n"
                "💡 Maslahat: Aniqroq va yorug' rasm yuboring.",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            _cleanup_file(image_path)
            await state.clear()
            await state.set_state(DocState.main_menu)
            return
        
        # === FAQAT MATN rejimi ===
        if selected_format == "Text":
            try:
                await status_msg.edit_text(
                    f"📋 Matn | 🔤 {selected_alphabet}\n\n"
                    f"✅ <i>Tayyor!</i>\n"
                    f"▓▓▓▓▓▓▓▓▓▓ 100%",
                    parse_mode="HTML"
                )
            except Exception:
                pass
            
            # Matnni xabar sifatida yuborish (4096 belgi limit)
            word_count = len(extracted_text.split())
            if len(extracted_text) <= 4000:
                await message.answer(
                    f"📋 <b>Ajratilgan matn</b> (~{word_count} so'z)\n"
                    f"━━━━━━━━━━━━━━━━━━━━\n\n"
                    f"{html_module.escape(extracted_text)}",
                    parse_mode="HTML"
                )
            else:
                # Uzun matnni bo'laklarga ajratish
                chunks = [extracted_text[i:i+3900] for i in range(0, len(extracted_text), 3900)]
                for idx, chunk in enumerate(chunks):
                    await message.answer(
                        f"📋 <b>Matn ({idx+1}/{len(chunks)})</b>\n\n"
                        f"{html_module.escape(chunk)}",
                        parse_mode="HTML"
                    )
            
            await message.answer("🏠 Davom etasizmi?", reply_markup=get_after_doc_keyboard())
            _cleanup_file(image_path)
            if cropped_path:
                _cleanup_file(cropped_path)
            await state.clear()
            await state.set_state(DocState.main_menu)
            return
        
        # 3-bosqich: Hujjat yaratilmoqda
        try:
            await status_msg.edit_text(
                f"{fmt_emoji} {fmt_label} | 🔤 {selected_alphabet}\n\n"
                f"📝 <i>Hujjat yaratilmoqda...</i>\n"
                f"▓▓▓▓▓▓▓░░░ 70%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        base_name = f"doc_{user_id}_{int(time.time())}"
        
        if selected_format == "PassPDF":
            output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
            await asyncio.to_thread(create_password_pdf, extracted_text, output_path, password, cropped_path)
            doc_file = FSInputFile(output_path, filename="Hujjat_parolli.pdf")
        elif selected_format == "PDF":
            output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
            await asyncio.to_thread(create_pdf_document, extracted_text, output_path, cropped_path)
            doc_file = FSInputFile(output_path, filename="Hujjat.pdf")
        else:
            output_path = os.path.join(TEMP_DIR, f"{base_name}.docx")
            await asyncio.to_thread(create_word_document, extracted_text, output_path, cropped_path)
            doc_file = FSInputFile(output_path, filename="Hujjat.docx")
        
        # 4-bosqich: Tayyor
        try:
            await status_msg.edit_text(
                f"{fmt_emoji} {fmt_label} | 🔤 {selected_alphabet}\n\n"
                f"✅ <i>Tayyor! Yuborilmoqda...</i>\n"
                f"▓▓▓▓▓▓▓▓▓▓ 100%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        word_count = len(extracted_text.split())
        
        # Sifat balli
        quality = await calculate_quality_score(extracted_text)
        stars = "⭐" * (quality["overall"] // 20) if quality["overall"] > 0 else "⭐⭐⭐"
        
        # Hujjat hash
        doc_hash = database.generate_doc_hash(extracted_text, user_id)
        database.save_doc_hash(doc_hash, user_id, fmt_label, word_count)
        
        caption = (
            f"✅ <b>Hujjat tayyor!</b>\n\n"
            f"📄 Format: <b>{fmt_label}</b>\n"
            f"🔤 Alifbo: <b>{selected_alphabet}</b>\n"
            f"📸 Rasmlar: <b>1</b> ta\n"
            f"📊 ~{word_count} so'z\n"
            f"📈 Sifat: {stars} ({quality['overall']}%)\n"
            f"🔐 ID: <code>{doc_hash}</code>"
        )
        if selected_format == "PassPDF":
            caption += f"\n🔒 Parol bilan himoyalangan"
        
        await message.answer_document(
            document=doc_file,
            caption=caption,
            parse_mode="HTML"
        )
        await message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
        await message.answer(
            "Hujjat sifati qanday bo'ldi?",
            reply_markup=get_feedback_keyboard(doc_hash)
        )
        
        _cleanup_file(image_path)
        if cropped_path:
            _cleanup_file(cropped_path)
        _cleanup_file(output_path)
        
    except Exception as e:
        logger.error(f"Xatolik: {e}", exc_info=True)
        error_msg = (
            "❌ <b>Xatolik yuz berdi</b>\n\n"
            "🔄 Iltimos, qayta urinib ko'ring.\n"
            "Muammo davom etsa /start bosing."
        )
        if ADMIN_ID and user_id == ADMIN_ID:
            error_msg += f"\n\n🔧 <code>{html_module.escape(str(e)[:500])}</code>"
        await message.answer(error_msg, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")
        _cleanup_file(image_path)
        if 'cropped_path' in locals() and cropped_path:
            _cleanup_file(cropped_path)
        if 'output_path' in locals() and output_path:
            _cleanup_file(output_path)
    
    finally:
        await state.clear()
        await state.set_state(DocState.main_menu)


# ==================== PAROL QABUL QILISH ====================

@dp.message(DocState.waiting_for_password, F.text)
async def password_handler(message: types.Message, state: FSMContext):
    """Parolni qabul qilish va hujjat yaratish."""
    if message.text.startswith('/'):
        return
    
    password = message.text.strip()
    if len(password) < 4:
        await message.answer(
            "⚠️ Parol kamida <b>4 belgi</b> bo'lishi kerak!",
            parse_mode="HTML"
        )
        return
    
    await _process_single_image(message, state, message.from_user.id, password=password)


# ==================== KO'P RASMLI HUJJAT ====================

async def _handle_multi_image(message: types.Message, state: FSMContext, image_path: str):
    """Ko'p rasmli hujjatga yangi rasm qo'shish."""
    data = await state.get_data()
    images = data.get("multi_images", [])
    target_count = data.get("multi_count", MAX_MULTI_IMAGES)
    
    images.append(image_path)
    await state.update_data(multi_images=images)
    
    current = len(images)
    
    # Agar kerakli songa yetgan bo'lsa — avtomatik davom etish
    if current >= target_count:
        # Limit tekshirish
        is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
        if not is_admin:
            has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
            if not has_limit:
                await message.answer(
                    "⚠️ <b>Bugungi limitingiz tugadi!</b>",
                    reply_markup=get_after_doc_keyboard(),
                    parse_mode="HTML"
                )
                for img in images:
                    _cleanup_file(img)
                await state.clear()
                await state.set_state(DocState.main_menu)
                return
        
        await message.answer(
            f"✅ <b>Barcha {target_count} ta rasm qabul qilindi!</b>\n\n"
            f"Qaysi formatda hujjat olmoqchisiz?",
            reply_markup=get_format_keyboard(),
            parse_mode="HTML"
        )
        await state.set_state(DocState.multi_format)
    else:
        remaining = target_count - current
        await message.answer(
            f"✅ Rasm #{current} qabul qilindi!\n\n"
            f"📸 Yuborilgan: <b>{current}</b> / {target_count}\n"
            f"📎 Yana <b>{remaining}</b> ta rasm yuboring.",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="main_menu")],
            ]),
            parse_mode="HTML"
        )


@dp.message(DocState.collecting_images, F.photo)
async def multi_collect_photo(message: types.Message, state: FSMContext):
    photo_file = await bot.get_file(message.photo[-1].file_id)
    image_path = os.path.join(TEMP_DIR, f"multi_{message.from_user.id}_{message.message_id}_{int(time.time())}.jpg")
    await bot.download_file(photo_file.file_path, image_path)
    await _handle_multi_image(message, state, image_path)


@dp.message(DocState.collecting_images, F.document)
async def multi_collect_document(message: types.Message, state: FSMContext):
    doc = message.document
    if not doc.mime_type or not doc.mime_type.startswith("image/"):
        await message.answer("⚠️ Faqat rasm fayllari qabul qilinadi.")
        return
    
    file = await bot.get_file(doc.file_id)
    ext = os.path.splitext(doc.file_name or "img.jpg")[1] or ".jpg"
    image_path = os.path.join(TEMP_DIR, f"multi_{message.from_user.id}_{message.message_id}_{int(time.time())}{ext}")
    await bot.download_file(file.file_path, image_path)
    await _handle_multi_image(message, state, image_path)

@dp.message(DocState.collecting_images)
async def multi_collect_fallback(message: types.Message):
    """Ko'p rasmli rejimda rasm bo'lmagan content uchun."""
    if message.text and message.text.startswith('/'):
        return
    await message.answer("📸 Faqat rasm yuboring!", parse_mode="HTML")


@dp.callback_query(DocState.multi_format, F.data.startswith("fmt_"))
async def multi_format_callback(callback: CallbackQuery, state: FSMContext):
    selected_format = callback.data.split("_", 1)[1]
    
    # Rasm sifatida (matn ajratilmaydi) — ko'p rasmli
    if selected_format in ("RawPDF", "RawWord"):
        await state.set_state('processing')
        data = await state.get_data()
        images = data.get("multi_images", [])
        user_id = callback.from_user.id
        total = len(images)
        
        if not images:
            await callback.message.edit_text("❌ Rasmlar topilmadi. /start bosing.")
            await state.clear()
            await callback.answer()
            return
        
        actual_format = "PDF" if selected_format == "RawPDF" else "Word"
        fmt_emoji = "📕" if actual_format == "PDF" else "📘"
        
        status_msg = await callback.message.edit_text(
            f"🖼 {actual_format} (Rasm sifatida)\n\n"
            f"⏳ <i>{total} ta rasm bilan hujjat yaratilmoqda...</i>\n"
            f"▓▓▓▓▓░░░░░ 50%",
            parse_mode="HTML"
        )
        await callback.answer()
        
        try:
            base_name = f"multi_raw_{user_id}_{int(time.time())}"
            if actual_format == "PDF":
                output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
                await asyncio.to_thread(create_multi_image_as_pdf, images, output_path)
                doc_file = FSInputFile(output_path, filename="Birlashtirilgan.pdf")
            else:
                output_path = os.path.join(TEMP_DIR, f"{base_name}.docx")
                await asyncio.to_thread(create_multi_image_as_word, images, output_path)
                doc_file = FSInputFile(output_path, filename="Birlashtirilgan.docx")
            
            try:
                await status_msg.edit_text(
                    f"🖼 {actual_format} (Rasm sifatida)\n\n"
                    f"✅ <i>{total} sahifali hujjat tayyor!</i>\n"
                    f"▓▓▓▓▓▓▓▓▓▓ 100%",
                    parse_mode="HTML"
                )
            except Exception:
                pass
            
            doc_hash = database.generate_doc_hash(f"mraw_{time.time()}", user_id)
            database.save_doc_hash(doc_hash, user_id, f"Ko'p Rasm -> {actual_format}", 0)
            
            await callback.message.answer_document(
                document=doc_file,
                caption=(
                    f"✅ <b>Hujjat tayyor!</b>\n\n"
                    f"📄 Format: <b>{actual_format}</b> (Rasm sifatida)\n"
                    f"📑 Sahifalar: <b>{total}</b>\n"
                    f"🔐 ID: <code>{doc_hash}</code>"
                ),
                parse_mode="HTML"
            )
            await callback.message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
            
            _cleanup_file(output_path)
            for img in images:
                _cleanup_file(img)
                
        except Exception as e:
            logger.error(f"Ko'p rasmli raw xatolik: {e}", exc_info=True)
            await callback.message.answer(
                "❌ <b>Xatolik yuz berdi</b>\n\n🔄 Qayta urinib ko'ring.",
                reply_markup=get_after_doc_keyboard(), parse_mode="HTML"
            )
            for img in images:
                _cleanup_file(img)
            if 'output_path' in locals() and output_path:
                _cleanup_file(output_path)
        finally:
            await state.clear()
            await state.set_state(DocState.main_menu)
        return
    
    # Oddiy matn ajratish oqimi
    await state.update_data(format=selected_format)
    
    fmt_map = {"PDF": "📕", "Word": "📘", "PassPDF": "🔒", "Text": "📋"}
    fmt_emoji = fmt_map.get(selected_format, "📄")
    fmt_label = "Parolli PDF" if selected_format == "PassPDF" else ("Faqat matn" if selected_format == "Text" else selected_format)
    
    await callback.message.edit_text(
        f"{fmt_emoji} Format: <b>{fmt_label}</b> ✅\n\nQaysi alifboda?",
        reply_markup=get_alphabet_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(DocState.multi_alphabet)
    await callback.answer()


@dp.callback_query(DocState.multi_alphabet, F.data.startswith("abc_"))
async def multi_alphabet_callback(callback: CallbackQuery, state: FSMContext):
    await state.set_state('processing')
    selected_alphabet = callback.data.split("_")[1]
    data = await state.get_data()
    images = data.get("multi_images", [])
    selected_format = data.get("format")
    total = len(images)
    
    await callback.message.edit_text(
        f"⏳ <i>{total} ta rasm tahlil qilinmoqda...</i>\n"
        f"▓░░░░░░░░░ 0/{total}",
        parse_mode="HTML"
    )
    await callback.answer()
    
    try:
        # Barcha rasmlarni tahlil qilish + progress
        texts = []
        cropped_images_list = []
        txt = None
        for i, img_path in enumerate(images):
            try:
                txt, cropped_path = await process_image_async(img_path, selected_alphabet)
                if txt and len(txt.strip()) >= 3:
                    texts.append(txt)
                    cropped_images_list.append(cropped_path)
            except Exception as e:
                logger.warning(f"Rasm #{i+1} tahlilida xatolik: {e}")
            
            # Progress yangilash
            done = i + 1
            pct = int((done / total) * 10)
            bar = "▓" * pct + "░" * (10 - pct)
            try:
                await callback.message.edit_text(
                    f"🔍 <i>Rasmlar tahlil qilinmoqda...</i>\n"
                    f"{bar} {done}/{total}\n\n"
                    f"📸 Rasm #{done} — {'✅' if (txt and len(txt.strip()) >= 3) else '⚠️'}",
                    parse_mode="HTML"
                )
            except Exception:
                pass
        
        if not texts:
            await callback.message.edit_text(
                "⚠️ <b>Rasmlardan matn topilmadi</b>\n\n"
                "💡 Aniqroq rasmlar yuboring.",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            for img in images:
                _cleanup_file(img)
            await state.clear()
            await state.set_state(DocState.main_menu)
            return
        
        # === FAQAT MATN rejimi (ko'p rasmli) ===
        if selected_format == "Text":
            try:
                await callback.message.edit_text(
                    f"📋 Matn | 🔤 {selected_alphabet}\n\n"
                    f"✅ <i>Tayyor!</i>\n"
                    f"▓▓▓▓▓▓▓▓▓▓ 100%",
                    parse_mode="HTML"
                )
            except Exception:
                pass
            
            total_words = sum(len(t.split()) for t in texts)
            full_text = "\n\n━━━━━━━━━━━━━━━━━━━━\n\n".join(texts)
            
            if len(full_text) <= 4000:
                await callback.message.answer(
                    f"📋 <b>Ajratilgan matn</b> ({len(texts)} sahifa, ~{total_words} so'z)\n"
                    f"━━━━━━━━━━━━━━━━━━━━\n\n"
                    f"{html_module.escape(full_text)}",
                    parse_mode="HTML"
                )
            else:
                for idx, text in enumerate(texts):
                    chunk = text[:3900]
                    await callback.message.answer(
                        f"📋 <b>Sahifa {idx+1}/{len(texts)}</b>\n\n"
                        f"{html_module.escape(chunk)}",
                        parse_mode="HTML"
                    )
            
            await callback.message.answer("🏠 Davom etasizmi?", reply_markup=get_after_doc_keyboard())
            for img in images:
                _cleanup_file(img)
            for img in cropped_images_list:
                if img: _cleanup_file(img)
            await state.clear()
            await state.set_state(DocState.main_menu)
            return
        
        # Hujjat yaratish
        try:
            await callback.message.edit_text(
                f"📝 <i>Hujjat yaratilmoqda...</i>\n"
                f"▓▓▓▓▓▓▓▓▓░ {total}/{total}\n\n"
                f"✅ {len(texts)} sahifa tayyor",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        base_name = f"multi_{callback.from_user.id}_{int(time.time())}"
        fmt_map = {"PDF": "📕", "Word": "📘", "PassPDF": "📕"}
        fmt_emoji = fmt_map.get(selected_format, "📄")
        fmt_label = selected_format
        if selected_format == "PassPDF":
            fmt_label = "PDF"
        
        if selected_format in ("PassPDF", "PDF"):
            output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
            await asyncio.to_thread(create_multi_image_pdf, texts, output_path, cropped_images_list)
            doc_file = FSInputFile(output_path, filename="Hujjat.pdf")
        else:
            output_path = os.path.join(TEMP_DIR, f"{base_name}.docx")
            await asyncio.to_thread(create_multi_image_word, texts, output_path, cropped_images_list)
            doc_file = FSInputFile(output_path, filename="Hujjat.docx")
        
        total_words = sum(len(t.split()) for t in texts)
        combined_text = "\n".join(texts)
        doc_hash = database.generate_doc_hash(combined_text, callback.from_user.id)
        database.save_doc_hash(doc_hash, callback.from_user.id, f"Ko'p rasmli {fmt_label}", total_words)
        
        await callback.message.answer_document(
            document=doc_file,
            caption=(
                f"✅ <b>Hujjat tayyor!</b>\n\n"
                f"📄 Format: <b>{fmt_label}</b>\n"
                f"🔤 Alifbo: <b>{selected_alphabet}</b>\n"
                f"📸 Rasmlar: <b>{len(texts)}</b> ta\n"
                f"📊 ~{total_words} so'z\n"
                f"🔐 ID: <code>{doc_hash}</code>"
            ),
            parse_mode="HTML"
        )
        await callback.message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
        await callback.message.answer(
            "Natija yoqdimi?",
            reply_markup=get_feedback_keyboard(doc_hash)
        )
        
        _cleanup_file(output_path)
        for img in images:
            _cleanup_file(img)
        for img in cropped_images_list:
            if img: _cleanup_file(img)
        
    except Exception as e:
        logger.error(f"Ko'p rasmli xatolik: {e}", exc_info=True)
        error_msg = (
            "❌ <b>Xatolik yuz berdi</b>\n\n"
            "🔄 Iltimos, qayta urinib ko'ring.\n"
            "Muammo davom etsa /start bosing."
        )
        if ADMIN_ID and callback.from_user.id == ADMIN_ID:
            error_msg += f"\n\n🔧 <code>{html_module.escape(str(e)[:500])}</code>"
        await callback.message.answer(error_msg, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")
        for img in images:
            _cleanup_file(img)
        if 'cropped_images_list' in locals():
            for img in cropped_images_list:
                if img: _cleanup_file(img)
        if 'output_path' in locals() and output_path:
            _cleanup_file(output_path)
    
    finally:
        await state.clear()
        await state.set_state(DocState.main_menu)


# ==================== KO'P RASMLI ASL PDF/WORD ====================

@dp.message(DocState.multi_raw_count, F.text)
async def multi_raw_count_handler(message: types.Message, state: FSMContext):
    if message.text.startswith('/'):
        return
    
    text = message.text.strip()
    if not text.isdigit():
        await message.answer("⚠️ Iltimos, faqat son kiriting (1-100).", parse_mode="HTML")
        return
    
    count = int(text)
    if count < 1 or count > MAX_MULTI_IMAGES:
        await message.answer(f"⚠️ Son 1 dan {MAX_MULTI_IMAGES} gacha bo'lishi kerak.", parse_mode="HTML")
        return
    
    await state.update_data(multi_raw_count=count, multi_raw_images=[])
    await message.answer(
        f"✅ <b>{count} ta rasm</b> kutilmoqda.\n\n"
        f"📸 Rasmlarni doimgidek bitta-bitta yuboring.\n"
        f"Yuborilgan: <b>0</b> / {count}",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.collecting_raw_images)

async def _handle_multi_raw_image(message: types.Message, state: FSMContext, image_path: str):
    data = await state.get_data()
    images = data.get("multi_raw_images", [])
    target_count = data.get("multi_raw_count", MAX_MULTI_IMAGES)
    
    images.append(image_path)
    await state.update_data(multi_raw_images=images)
    
    current = len(images)
    
    if current >= target_count:
        is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
        if not is_admin:
            has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
            if not has_limit:
                await message.answer(
                    "⚠️ <b>Bugungi limitingiz tugadi!</b>",
                    reply_markup=get_after_doc_keyboard(),
                    parse_mode="HTML"
                )
                for img in images:
                    _cleanup_file(img)
                await state.clear()
                await state.set_state(DocState.main_menu)
                return
        
        await message.answer(
            f"✅ <b>Barcha {target_count} ta rasm qabul qilindi!</b>\n\n"
            f"Qaysi formatda hujjat olmoqchisiz?",
            reply_markup=get_raw_format_keyboard(),
            parse_mode="HTML"
        )
        await state.set_state(DocState.multi_raw_format)
    else:
        remaining = target_count - current
        await message.answer(
            f"✅ Rasm #{current} qabul qilindi!\n\n"
            f"📸 Yuborilgan: <b>{current}</b> / {target_count}\n"
            f"📎 Yana <b>{remaining}</b> ta rasm yuboring.",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="main_menu")],
            ]),
            parse_mode="HTML"
        )

@dp.message(DocState.collecting_raw_images, F.photo)
async def multi_raw_collect_photo(message: types.Message, state: FSMContext):
    photo_file = await bot.get_file(message.photo[-1].file_id)
    image_path = os.path.join(TEMP_DIR, f"mraw_{message.from_user.id}_{message.message_id}_{int(time.time())}.jpg")
    await bot.download_file(photo_file.file_path, image_path)
    await _handle_multi_raw_image(message, state, image_path)

@dp.message(DocState.collecting_raw_images, F.document)
async def multi_raw_collect_document(message: types.Message, state: FSMContext):
    doc = message.document
    if not doc.mime_type or not doc.mime_type.startswith("image/"):
        await message.answer("⚠️ Faqat rasm fayllari qabul qilinadi.")
        return
    
    file = await bot.get_file(doc.file_id)
    ext = os.path.splitext(doc.file_name or "img.jpg")[1] or ".jpg"
    image_path = os.path.join(TEMP_DIR, f"mraw_{message.from_user.id}_{message.message_id}_{int(time.time())}{ext}")
    await bot.download_file(file.file_path, image_path)
    await _handle_multi_raw_image(message, state, image_path)

@dp.message(DocState.collecting_raw_images)
async def multi_raw_collect_fallback(message: types.Message):
    if message.text and message.text.startswith('/'):
        return
    await message.answer("📸 Faqat rasm yuboring!", parse_mode="HTML")

@dp.callback_query(DocState.multi_raw_format, F.data.startswith("rawfmt_"))
async def multi_raw_format_callback(callback: CallbackQuery, state: FSMContext):
    await state.set_state('processing')
    selected_format = callback.data.split("_", 1)[1]
    data = await state.get_data()
    images = data.get("multi_raw_images", [])
    user_id = callback.from_user.id
    total = len(images)
    
    if not images:
        await callback.message.edit_text("❌ Rasmlar topilmadi. /start bosing.")
        await state.clear()
        await callback.answer()
        return
    
    fmt_emoji = "📕" if selected_format == "PDF" else "📘"
    
    status_msg = await callback.message.edit_text(
        f"{fmt_emoji} {selected_format} (Asl rasmlar)\n\n"
        f"⏳ <i>{total} ta rasm bilan hujjat yaratilmoqda...</i>\n"
        f"▓▓▓▓▓░░░░░ 50%",
        parse_mode="HTML"
    )
    await callback.answer()
    
    try:
        base_name = f"multi_raw_{user_id}_{int(time.time())}"
        
        if selected_format == "PDF":
            output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
            await asyncio.to_thread(create_multi_image_as_pdf, images, output_path)
            doc_file = FSInputFile(output_path, filename="Birlashtirilgan.pdf")
        else:
            output_path = os.path.join(TEMP_DIR, f"{base_name}.docx")
            await asyncio.to_thread(create_multi_image_as_word, images, output_path)
            doc_file = FSInputFile(output_path, filename="Birlashtirilgan.docx")
        
        try:
            await status_msg.edit_text(
                f"{fmt_emoji} {selected_format} (Asl rasmlar)\n\n"
                f"✅ <i>{total} sahifali hujjat tayyor! Yuborilmoqda...</i>\n"
                f"▓▓▓▓▓▓▓▓▓▓ 100%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        doc_hash = database.generate_doc_hash(f"mraw_{time.time()}", user_id)
        database.save_doc_hash(doc_hash, user_id, f"Ko'p Asl Rasm -> {selected_format}", 0)
        
        caption = (
            f"✅ <b>Hujjat tayyor!</b>\n\n"
            f"📄 Format: <b>{selected_format}</b> (Asl rasm)\n"
            f"📑 Sahifalar: <b>{total}</b>\n"
            f"🔐 ID: <code>{doc_hash}</code>"
        )
        
        await callback.message.answer_document(
            document=doc_file,
            caption=caption,
            parse_mode="HTML"
        )
        await callback.message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
        
        _cleanup_file(output_path)
        for img in images:
            _cleanup_file(img)
            
    except Exception as e:
        logger.error(f"Ko'p asl rasmli xatolik: {e}", exc_info=True)
        error_msg = "❌ <b>Xatolik yuz berdi</b>\n\n🔄 Iltimos, qayta urinib ko'ring."
        await callback.message.answer(error_msg, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")
        for img in images:
            _cleanup_file(img)
        if 'output_path' in locals() and output_path:
            _cleanup_file(output_path)
    finally:
        await state.clear()
        await state.set_state(DocState.main_menu)


# ==================== /help ====================

HELP_TEXT = (
    "❓ <b>Qanday ishlaydi?</b>\n"
    "━━━━━━━━━━━━━━━━━━━━\n\n"
    "Bu bot rasmlardan matnni aqlli sun'iy\n"
    "intellekt yordamida ajratib, tayyor PDF\n"
    "yoki Word hujjat qilib beradi.\n\n"
    "━━━━━━━━━━━━━━━━━━━━\n\n"
    "📸 <b>Rasmni hujjatga aylantirish</b>\n"
    "Matnli rasmni yuboring (daftar, kitob..),\n"
    "format tanlang — tayyor PDF/Word oling!\n\n"
    "📚 <b>Ko'p rasm → 1 hujjat</b>\n"
    "100 tagacha rasmni bitta hujjatga\n"
    "birlashtirish mumkin.\n\n"
    "🎤 <b>Ovozni hujjatga aylantirish</b>\n"
    "Ovozli xabar yuboring — gapirgan\n"
    "matnni PDF hujjat qilib beradi!\n\n"
    "🌐 <b>Rasmni tarjima qilib hujjat</b>\n"
    "Rasmdan matnni ajratib, boshqa tilga\n"
    "tarjima qilib PDF qilib beradi.\n"
    "(🇺🇿↔🇷🇺 O'zbek↔Rus, 🇺🇿↔🇬🇧 O'zbek↔Ingliz)\n\n"
    "🔄 <b>Lotin ↔ Кирилл o'girish</b>\n"
    "Matn yuboring — avtomatik boshqa\n"
    "alifboga o'girib beradi.\n\n"
    "━━━━━━━━━━━━━━━━━━━━\n"
    "💡 <b>Foydali ma'lumot:</b>\n\n"
    "✍️ Qo'l yozuvini ham aniqlaydi\n"
    "📝 Imlo xatolari avtomatik tuzatiladi\n"
    "🎁 Har kuni <b>2 ta bepul</b> foydalanish\n"
    "🔥 Har kuni kirganingizda +bonus olasiz\n"
    "📈 Har hujjatga sifat balli beriladi"
)

def get_help_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📸 Rasm → Hujjat misoli", callback_data="help_example_image")],
        [InlineKeyboardButton(text="🏠 Bosh menyu", callback_data="main_menu")],
    ])

@dp.callback_query(F.data == "account_help_menu")
async def account_help_menu_callback(callback: CallbackQuery, state: FSMContext):
    uid = callback.from_user.id
    is_admin = (ADMIN_ID and uid == ADMIN_ID)
    text = (
        "⚙️ <b>Hisobim & Yordam</b>\n"
        "━━━━━━━━━━━━━━━━━━━━\n\n"
        "Shaxsiy ma'lumotlaringizni ko'rish yoki "
        "botdan qanday foydalanish qoidalari bilan tanishish uchun "
        "quyidagi bo'limlardan birini tanlang."
    )
    await callback.message.edit_text(
        text,
        reply_markup=get_account_help_keyboard(is_admin),
        parse_mode="HTML"
    )
    await callback.answer()

@dp.message(Command("help"), StateFilter("*"))
async def help_cmd(message: types.Message, state: FSMContext):
    await state.clear()
    await state.set_state(DocState.main_menu)
    await message.answer(
        HELP_TEXT,
        reply_markup=get_help_keyboard(),
        parse_mode="HTML"
    )

@dp.callback_query(F.data == "help_menu")
async def help_menu_callback(callback: CallbackQuery):
    kb = get_help_keyboard()
    # Orqaga tugmasini main_menu o'rniga account_help_menu qilamiz (agar iloji bo'lsa)
    # Biz to'g'ridan-to'g'ri yangi klaviatura yozamiz
    new_kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📸 Rasmni PDF/Word (misol)", callback_data="help_example_image")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="account_help_menu")]
    ])
    await callback.message.edit_text(
        HELP_TEXT,
        reply_markup=new_kb,
        parse_mode="HTML"
    )
    await callback.answer()

@dp.callback_query(F.data == "help_example_image")
async def help_example_image_callback(callback: CallbackQuery):
    example = (
        "📸 <b>Misol: Rasmni hujjatga aylantirish</b>\n"
        "━━━━━━━━━━━━━━━━━━━━\n\n"
        "<b>1-qadam:</b> Bosh menyuda\n"
        "   \"📸 Rasmni hujjatga aylantir\" bosing\n\n"
        "<b>2-qadam:</b> Matn yozilgan rasmni yuboring\n"
        "   <i>(daftar, kitob, konspekt rasmi)</i>\n\n"
        "<b>3-qadam:</b> Format tanlang:\n"
        "   📕 PDF — oddiy PDF fayl\n"
        "   📘 Word — tahrirlash mumkin\n"
        "   🔒 Parolli PDF — himoyalangan\n"
        "   📋 Faqat matn — nusxalash uchun\n\n"
        "<b>4-qadam:</b> Alifbo tanlang:\n"
        "   🔤 Lotin yoki 🔤 Кирилл\n\n"
        "✅ <b>Tayyor!</b> PDF/Word fayl keladi!\n\n"
        "💡 <i>Rasm qanchalik sifatli bo'lsa,\n"
        "natija shunchalik yaxshi chiqadi.</i>"
    )
    await callback.message.edit_text(
        example,
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="help_menu")],
            [InlineKeyboardButton(text="📸 Boshlash", callback_data="mode_image")],
            [InlineKeyboardButton(text="🏠 Bosh menyu", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await callback.answer()


# ==================== /myaccount ====================

@dp.message(Command("myaccount"), StateFilter("*"))
async def myaccount_cmd(message: types.Message, state: FSMContext):
    await state.clear()
    await state.set_state(DocState.main_menu)
    await _send_account_info(message)

@dp.callback_query(F.data == "myaccount")
async def myaccount_callback(callback: CallbackQuery):
    uid = callback.from_user.id
    is_admin = (ADMIN_ID and uid == ADMIN_ID)
    balance = "♾️" if is_admin else database.get_user_balance(uid)
    streak = database.get_streak(uid)
    streak_emoji = _get_streak_emoji(streak)
    rank = database.get_user_rank(uid)
    
    text = (
        f"📊 <b>Sizning Hisobingiz</b>\n"
        f"━━━━━━━━━━━━━━━━━━━━\n\n"
        f"🆔 ID: <code>{uid}</code>\n"
        f"🎯 Qoldiq: <b>{balance}</b> ta\n"
        f"📅 Kunlik tekin: 2 ta\n"
        f"{streak_emoji} Streak: <b>{streak}</b> kun\n"
        f"🏆 Reyting: #{rank}\n"
    )
    
    if streak >= 7:
        next_bonus = 7 - (streak % 7)
        text += f"\n🎁 Keyingi bonus: <b>{next_bonus}</b> kunga"
    
    keyboard = []
    if is_admin:
        # Reyting faqat adminga ko'rinadi
        pass
    keyboard.append([InlineKeyboardButton(text="⬅️ Orqaga", callback_data="account_help_menu")])
    
    try:
        await callback.message.edit_text(
            text,
            reply_markup=InlineKeyboardMarkup(inline_keyboard=keyboard),
            parse_mode="HTML"
        )
    except Exception:
        await callback.message.answer(text, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")
    await callback.answer()

async def _send_account_info(message: types.Message, user_id: int = None):
    uid = user_id or message.from_user.id
    is_admin = (ADMIN_ID and uid == ADMIN_ID)
    balance = "♾️" if is_admin else database.get_user_balance(uid)
    streak = database.get_streak(uid)
    streak_emoji = _get_streak_emoji(streak)
    rank = database.get_user_rank(uid)
    
    text = (
        f"📊 <b>Sizning Hisobingiz</b>\n"
        f"━━━━━━━━━━━━━━━━━━━━\n\n"
        f"🆔 ID: <code>{uid}</code>\n"
        f"🎯 Qoldiq: <b>{balance}</b> ta\n"
        f"📅 Kunlik tekin: 2 ta\n"
        f"{streak_emoji} Streak: <b>{streak}</b> kun\n"
        f"🏆 Reyting: #{rank}\n"
    )
    
    await message.answer(text, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")


# ==================== 🎤 OVOZLI XABAR ====================

@dp.callback_query(F.data == "mode_voice")
async def mode_voice(callback: CallbackQuery, state: FSMContext):
    await callback.message.edit_text(
        "🎤 <b>Ovoz → Hujjat</b>\n\n"
        "Ovozli xabar yuboring — AI matnni ajratib,\n"
        "PDF hujjat qilib beradi!\n\n"
        "🎙️ <i>Ovozli xabar yuboring:</i>",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.voice_select_alphabet)
    await callback.answer()


@dp.message(DocState.voice_select_alphabet, F.voice)
async def voice_received(message: types.Message, state: FSMContext):
    """Ovozli xabar qabul qilish."""
    is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
    if not is_admin:
        has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
        if not has_limit:
            await message.answer(
                "⚠️ <b>Bugungi limitingiz tugadi!</b>",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            return
    
    voice_file = await bot.get_file(message.voice.file_id)
    voice_path = os.path.join(TEMP_DIR, f"voice_{message.from_user.id}_{int(time.time())}.ogg")
    await bot.download_file(voice_file.file_path, voice_path)
    
    await state.update_data(voice_path=voice_path)
    
    await message.answer(
        "✅ Ovoz qabul qilindi!\n\n"
        "Qaysi alifboda yozilsin?",
        reply_markup=get_alphabet_keyboard(),
        parse_mode="HTML"
    )
    await state.set_state(DocState.voice_select_format)


@dp.message(DocState.voice_select_alphabet)
async def voice_not_voice(message: types.Message):
    if message.text and message.text.startswith('/'):
        return
    await message.answer("🎤 Iltimos, ovozli xabar yuboring!", parse_mode="HTML")


@dp.callback_query(DocState.voice_select_format, F.data.startswith("abc_"))
async def voice_alphabet_callback(callback: CallbackQuery, state: FSMContext):
    await state.set_state('processing')
    selected_alphabet = callback.data.split("_")[1]
    data = await state.get_data()
    voice_path = data.get("voice_path")
    
    if not voice_path or not os.path.exists(voice_path):
        await callback.message.edit_text("❌ Ovoz topilmadi. /start bosing.")
        await state.clear()
        await callback.answer()
        return
    
    status_msg = await callback.message.edit_text(
        "🎤 <b>Ovoz → Hujjat</b>\n\n"
        "⏳ <i>AI ovozni tahlil qilmoqda...</i>\n"
        "▓▓░░░░░░░░ 20%",
        parse_mode="HTML"
    )
    await callback.answer()
    
    try:
        # AI tahlil
        try:
            await callback.message.edit_text(
                "🎤 <b>Ovoz → Hujjat</b>\n\n"
                "🔍 <i>Matn aniqlanmoqda...</i>\n"
                "▓▓▓▓▓░░░░░ 50%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        extracted_text = await process_voice_async(voice_path, selected_alphabet)
        
        if not extracted_text or len(extracted_text.strip()) < 3:
            await callback.message.edit_text(
                "⚠️ <b>Ovozda matn topilmadi</b>\n\n"
                "💡 Aniqroq gapiring.",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            _cleanup_file(voice_path)
            await state.clear()
            await state.set_state(DocState.main_menu)
            return
        
        # PDF yaratish
        try:
            await callback.message.edit_text(
                "🎤 <b>Ovoz → Hujjat</b>\n\n"
                "📝 <i>PDF yaratilmoqda...</i>\n"
                "▓▓▓▓▓▓▓▓░░ 80%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        base_name = f"voice_{callback.from_user.id}_{int(time.time())}"
        output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
        await asyncio.to_thread(create_pdf_document, extracted_text, output_path)
        doc_file = FSInputFile(output_path, filename="Ovoz_hujjat.pdf")
        
        quality = await calculate_quality_score(extracted_text)
        stars = "⭐" * (quality["overall"] // 20) if quality["overall"] > 0 else "⭐⭐⭐"
        doc_hash = database.generate_doc_hash(extracted_text, callback.from_user.id)
        database.save_doc_hash(doc_hash, callback.from_user.id, "Ovoz→PDF", len(extracted_text.split()))
        
        try:
            await callback.message.edit_text(
                "🎤 <b>Ovoz → Hujjat</b>\n\n"
                "✅ <i>Tayyor!</i>\n"
                "▓▓▓▓▓▓▓▓▓▓ 100%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        word_count = len(extracted_text.split())
        await callback.message.answer_document(
            document=doc_file,
            caption=(
                f"✅ <b>Ovozdan hujjat tayyor!</b>\n\n"
                f"🎤 Manba: Ovozli xabar\n"
                f"🔤 Alifbo: <b>{selected_alphabet}</b>\n"
                f"📊 ~{word_count} so'z\n"
                f"📈 Sifat: {stars} ({quality['overall']}%)\n"
                f"🔐 ID: <code>{doc_hash}</code>"
            ),
            parse_mode="HTML"
        )
        await callback.message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
        await callback.message.answer(
            "Sifat qanday bo'ldi?",
            reply_markup=get_feedback_keyboard(doc_hash)
        )
        
        _cleanup_file(voice_path)
        _cleanup_file(output_path)
        
    except Exception as e:
        logger.error(f"Ovoz xatolik: {e}", exc_info=True)
        error_msg = "❌ <b>Xatolik yuz berdi</b>\n\n🔄 Qayta urinib ko'ring."
        if ADMIN_ID and callback.from_user.id == ADMIN_ID:
            error_msg += f"\n\n🔧 <code>{html_module.escape(str(e)[:500])}</code>"
        await callback.message.answer(error_msg, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")
        _cleanup_file(voice_path)
        if 'output_path' in locals() and output_path:
            _cleanup_file(output_path)
    
    finally:
        await state.clear()
        await state.set_state(DocState.main_menu)


# ==================== 🏆 REYTING ====================

@dp.callback_query(F.data == "leaderboard")
async def leaderboard_callback(callback: CallbackQuery):
    leaders = database.get_leaderboard(10)
    uid = callback.from_user.id
    rank = database.get_user_rank(uid)
    
    text = "🏆 <b>TOP Foydalanuvchilar</b>\n━━━━━━━━━━━━━━━━━━━━\n\n"
    
    medals = ["🥇", "🥈", "🥉"]
    for i, u in enumerate(leaders):
        medal = medals[i] if i < 3 else f"{i+1}."
        streak_e = _get_streak_emoji(u.get('streak_days', 0))
        name = u['first_name'][:15] if u['first_name'] else "—"
        is_you = " ← Siz" if u['user_id'] == uid else ""
        text += f"{medal} <b>{name}</b> — {u['total_docs_created']} ta hujjat {streak_e}{is_you}\n"
    
    text += f"\n━━━━━━━━━━━━━━━━━━━━\n📍 Sizning o'rningiz: <b>#{rank}</b>"
    
    try:
        await callback.message.edit_text(
            text,
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="🏠 Bosh menyu", callback_data="main_menu")],
            ]),
            parse_mode="HTML"
        )
    except Exception:
        await callback.message.answer(text, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")
    await callback.answer()


# ==================== 👍/👎 FEEDBACK ====================

@dp.callback_query(F.data.startswith("fb_"))
async def feedback_callback(callback: CallbackQuery):
    parts = callback.data.split("_", 2)
    if len(parts) < 3:
        await callback.answer("❌ Xato")
        return
    
    rating_type = parts[1]  # good or bad
    doc_hash = parts[2]
    rating = 5 if rating_type == "good" else 1
    
    database.save_feedback(callback.from_user.id, doc_hash, rating)
    
    if rating_type == "good":
        response = "👍 Rahmat! Fikringiz biz uchun muhim!"
    else:
        response = "👎 Tushundim. Biz sifatni yaxshilashga harakat qilamiz!"
    
    try:
        await callback.message.edit_text(
            f"{response}\n\n🏠 Bosh menyuga qaytish uchun bosing:",
            reply_markup=get_after_doc_keyboard(),
            parse_mode="HTML"
        )
    except Exception:
        pass
    await callback.answer(response)


# ==================== 🔐 VERIFY HUJJAT ====================

@dp.message(Command("verify"), StateFilter("*"))
async def verify_cmd(message: types.Message, state: FSMContext):
    """Hujjat haqiqiyligini tekshirish."""
    await state.clear()
    await state.set_state(DocState.main_menu)
    args = message.text.replace("/verify", "", 1).strip()
    
    if not args:
        await message.answer(
            "🔐 <b>Hujjat tekshirish</b>\n\n"
            "Foydalanish: <code>/verify HUJJAT_ID</code>\n\n"
            "💡 Hujjat ID hujjat tayyor bo'lganda beriladi.",
            parse_mode="HTML"
        )
        return
    
    result = database.verify_doc_hash(args.strip())
    
    if result:
        await message.answer(
            f"✅ <b>Hujjat tasdiqlandi!</b>\n"
            f"━━━━━━━━━━━━━━━━━━━━\n\n"
            f"🔐 ID: <code>{result['doc_hash']}</code>\n"
            f"📄 Turi: <b>{result['doc_type']}</b>\n"
            f"📊 ~{result['word_count']} so'z\n"
            f"📅 Yaratilgan: <b>{result['created_at']}</b>\n\n"
            f"✅ Bu hujjat AI Hujjat Bot orqali yaratilgan.",
            reply_markup=get_after_doc_keyboard(),
            parse_mode="HTML"
        )
    else:
        await message.answer(
            "❌ <b>Hujjat topilmadi</b>\n\n"
            "Bu ID bo'yicha hujjat mavjud emas.\n"
            "ID to'g'ri kiritilganligini tekshiring.",
            reply_markup=get_after_doc_keyboard(),
            parse_mode="HTML"
        )


# ==================== 📋 SMART CLIPBOARD ====================

@dp.message(F.forward_from | F.forward_from_chat)
async def smart_clipboard_handler(message: types.Message, state: FSMContext):
    """Forward qilingan matnni hujjatga aylantirish."""
    current_state = await state.get_state()
    if current_state and current_state != DocState.main_menu.state:
        return
    await state.set_state('processing')
    
    if not message.text or len(message.text.strip()) < 10:
        await state.clear()
        await state.set_state(DocState.main_menu)
        return
    
    is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
    if not is_admin:
        has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
        if not has_limit:
            await message.answer(
                "⚠️ <b>Bugungi limitingiz tugadi!</b>",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            return
    
    text = message.text.strip()
    word_count = len(text.split())
    
    status_msg = await message.answer("⏳ Ishlanmoqda, kuting..."
        "📝 <i>Forward matn → PDF yaratilmoqda...</i>\n"
        "▓▓▓▓▓░░░░░ 50%",
        parse_mode="HTML"
    )
    
    try:
        base_name = f"clip_{message.from_user.id}_{int(time.time())}"
        output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
        await asyncio.to_thread(create_pdf_document, text, output_path)
        doc_file = FSInputFile(output_path, filename="Clipboard.pdf")
        
        doc_hash = database.generate_doc_hash(text, message.from_user.id)
        database.save_doc_hash(doc_hash, message.from_user.id, "Clipboard→PDF", word_count)
        
        try:
            await status_msg.edit_text(
                "📋 <b>Smart Clipboard</b>\n\n"
                "✅ <i>Tayyor!</i>\n"
                "▓▓▓▓▓▓▓▓▓▓ 100%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        await message.answer_document(
            document=doc_file,
            caption=(
                f"✅ <b>Forward → PDF tayyor!</b>\n\n"
                f"📊 ~{word_count} so'z\n"
                f"🔐 ID: <code>{doc_hash}</code>"
            ),
            parse_mode="HTML"
        )
        await message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
        
        _cleanup_file(output_path)
        
    except Exception as e:
        logger.error(f"Clipboard xatolik: {e}", exc_info=True)
        await message.answer("❌ Xatolik yuz berdi.", reply_markup=get_after_doc_keyboard())


# ==================== ADMIN ====================

@dp.message(Command("admin"), StateFilter("*"))
async def admin_panel(message: types.Message):
    if ADMIN_ID and message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Faqat admin uchun.")
        return
    
    stats = database.get_user_stats()
    fb_stats = database.get_feedback_stats()
    users = database.get_all_users_list(30)
    
    text = (
        f"📊 <b>ADMIN PANEL</b>\n\n"
        f"👥 Jami: <b>{stats['total_users']}</b>\n"
        f"📈 Bugun faol: <b>{stats['today_active']}</b>\n"
        f"📄 Jami hujjatlar: <b>{stats['total_docs']}</b>\n"
        f"⭐ Feedback: {fb_stats['avg_rating']}/5 ({fb_stats['total']} ta)\n\n"
        f"━━━━━━━━━━━━━━━━━━━━\n"
    )
    
    for i, u in enumerate(users[:20], 1):
        un = f"@{u['username']}" if u['username'] else "—"
        text += f"{i}. <code>{u['user_id']}</code> | {u['first_name']} | {un} | D:{u['total_docs_created']}\n"
    
    await message.answer(text, parse_mode="HTML")


@dp.message(Command("broadcast"), StateFilter("*"))
async def broadcast_cmd(message: types.Message):
    if ADMIN_ID and message.from_user.id != ADMIN_ID:
        await message.answer("⛔ Faqat admin uchun.")
        return
    
    text_to_send = message.text.replace("/broadcast", "", 1).strip()
    if not text_to_send:
        await message.answer("Foydalanish: /broadcast Xabar matni")
        return
    
    user_ids = database.get_all_user_ids()
    sent = failed = 0
    status_msg = await message.answer(f"📤 0/{len(user_ids)}")
    
    for uid in user_ids:
        try:
            await bot.send_message(uid, text_to_send, parse_mode="HTML")
            sent += 1
        except Exception:
            failed += 1
        if (sent + failed) % 20 == 0:
            await asyncio.sleep(1.5)
            try:
                await status_msg.edit_text(f"📤 {sent + failed}/{len(user_ids)}")
            except Exception:
                pass
    
    await status_msg.edit_text(
        f"✅ <b>Broadcast tugadi!</b>\n📤 {sent} | ❌ {failed} | 👥 {len(user_ids)}",
        parse_mode="HTML"
    )


# ==================== PDF <-> WORD KONVERTOR ====================

@dp.callback_query(F.data == "mode_pdf_to_word")
async def mode_pdf_to_word(callback: CallbackQuery, state: FSMContext):
    await state.clear()
    await callback.message.edit_text(
        "📕 <b>PDF → Word</b>\n"
        "━━━━━━━━━━━━━━━━━━━━\n\n"
        "PDF faylni menga yuboring —\n"
        "men uni Word (.docx) formatiga\n"
        "aylantirib beraman!\n\n"
        "👇 <b>PDF faylni hozir yuboring:</b>",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.waiting_for_pdf_file)
    await callback.answer()


@dp.callback_query(F.data == "mode_word_to_pdf")
async def mode_word_to_pdf(callback: CallbackQuery, state: FSMContext):
    await state.clear()
    await callback.message.edit_text(
        "📘 <b>Word → PDF</b>\n"
        "━━━━━━━━━━━━━━━━━━━━\n\n"
        "Word (.docx) faylni menga yuboring —\n"
        "men uni PDF formatiga\n"
        "aylantirib beraman!\n\n"
        "👇 <b>Word faylni hozir yuboring:</b>",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.waiting_for_word_file)
    await callback.answer()


@dp.message(DocState.waiting_for_pdf_file, F.document)
async def receive_pdf_file(message: types.Message, state: FSMContext):
    """PDF faylni qabul qilish."""
    doc = message.document
    if not doc.file_name or not doc.file_name.lower().endswith('.pdf'):
        await message.answer(
            "⚠️ Bu PDF fayl emas!\n"
            "Iltimos, <b>.pdf</b> formatdagi fayl yuboring.",
            parse_mode="HTML"
        )
        return
    
    is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
    if not is_admin:
        has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
        if not has_limit:
            await message.answer(
                "⚠️ <b>Bugungi limitingiz tugadi!</b>\n💡 Ertaga yana 2 ta tekin beriladi.",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            return
    
    file = await bot.get_file(doc.file_id)
    file_path = os.path.join(TEMP_DIR, f"conv_{message.from_user.id}_{int(time.time())}.pdf")
    await bot.download_file(file.file_path, file_path)
    
    await state.update_data(convert_file_path=file_path, convert_direction="pdf_to_word", original_filename=doc.file_name)
    await message.answer(
        "✅ PDF fayl qabul qilindi!\n\n"
        "🔤 Qaysi alifboda chiqarish kerak?\n"
        "💡 <i>(Agar PDF skanerlangan rasm bo'lsa, bot uni AI orqali yozmaga aylantiradi)</i>",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [
                InlineKeyboardButton(text="🔤 Lotin", callback_data="convabc_Lotin"),
                InlineKeyboardButton(text="🔤 Кирилл", callback_data="convabc_Kirill")
            ],
            [InlineKeyboardButton(text="🔤 O'zgartirmasdan", callback_data="convabc_Original")],
            [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.convert_select_alphabet)


@dp.message(DocState.waiting_for_pdf_file)
async def pdf_file_fallback(message: types.Message):
    if message.text and message.text.startswith('/'):
        return
    await message.answer("📄 Iltimos, PDF faylni yuboring!", parse_mode="HTML")


@dp.message(DocState.waiting_for_word_file, F.document)
async def receive_word_file(message: types.Message, state: FSMContext):
    """Word faylni qabul qilish."""
    doc = message.document
    fname = (doc.file_name or "").lower()
    
    if fname.endswith('.doc') and not fname.endswith('.docx'):
        await message.answer(
            "⚠️ <b>Eski .doc formati qo'llab-quvvatlanmaydi!</b>\n\n"
            "Iltimos, faylni <b>.docx</b> formatda saqlang:\n"
            "📌 Word → Fayl → Saqlash → .docx tanlang\n\n"
            "Keyin qayta yuboring.",
            parse_mode="HTML"
        )
        return
    
    if not fname.endswith('.docx'):
        await message.answer(
            "⚠️ Bu Word fayl emas!\n"
            "Iltimos, <b>.docx</b> formatdagi fayl yuboring.",
            parse_mode="HTML"
        )
        return
    
    is_admin = (ADMIN_ID and message.from_user.id == ADMIN_ID)
    if not is_admin:
        has_limit, remaining = database.check_and_deduct_limit(message.from_user.id)
        if not has_limit:
            await message.answer(
                "⚡ <b>Bugungi limitingiz tugadi!</b>\n"
                "⏰ Ertaga soat 00:00 da yangilanadi\n"
                "💎 Premium olish → /premium",
                reply_markup=get_after_doc_keyboard(),
                parse_mode="HTML"
            )
            return
    
    file = await bot.get_file(doc.file_id)
    file_path = os.path.join(TEMP_DIR, f"conv_{message.from_user.id}_{int(time.time())}.docx")
    await bot.download_file(file.file_path, file_path)
    
    await state.update_data(convert_file_path=file_path, convert_direction="word_to_pdf", original_filename=doc.file_name)
    await message.answer(
        "✅ Word fayl qabul qilindi!\n\n"
        "🔤 Qaysi alifboda chiqarish kerak?",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [
                InlineKeyboardButton(text="🔤 Lotin", callback_data="convabc_Lotin"),
                InlineKeyboardButton(text="🔤 Кирилл", callback_data="convabc_Kirill")
            ],
            [InlineKeyboardButton(text="🔤 O'zgartirmasdan", callback_data="convabc_Original")],
            [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="main_menu")],
        ]),
        parse_mode="HTML"
    )
    await state.set_state(DocState.convert_select_alphabet)


@dp.message(DocState.waiting_for_word_file)
async def word_file_fallback(message: types.Message):
    if message.text and message.text.startswith('/'):
        return
    await message.answer("📄 Iltimos, Word faylni yuboring!", parse_mode="HTML")


@dp.callback_query(DocState.convert_select_alphabet, F.data.startswith("convabc_"))
async def convert_alphabet_callback(callback: CallbackQuery, state: FSMContext):
    """PDF<->Word konvertatsiya — alifbo tanlangandan keyin."""
    await state.set_state('processing')
    selected_alphabet = callback.data.split("_", 1)[1]
    data = await state.get_data()
    file_path = data.get("convert_file_path")
    direction = data.get("convert_direction")
    original_name = data.get("original_filename", "Fayl")
    user_id = callback.from_user.id
    
    if not file_path or not os.path.exists(file_path):
        await callback.message.edit_text("❌ Fayl topilmadi. /start bosing.")
        await state.clear()
        await callback.answer()
        return
    
    is_pdf_to_word = (direction == "pdf_to_word")
    direction_label = "PDF → Word" if is_pdf_to_word else "Word → PDF"
    direction_emoji = "📕→📘" if is_pdf_to_word else "📘→📕"
    
    status_msg = await callback.message.edit_text(
        f"{direction_emoji} <b>{direction_label}</b>\n\n"
        f"⏳ <i>Konvertatsiya qilinmoqda...</i>\n"
        f"▓▓▓░░░░░░░ 30%",
        parse_mode="HTML"
    )
    await callback.answer()
    
    output_path = None
    
    try:
        base_name = f"conv_{user_id}_{int(time.time())}"
        
        if is_pdf_to_word:
            import fitz
            output_path = os.path.join(TEMP_DIR, f"{base_name}.docx")
            
            doc = fitz.open(file_path)
            logger.info(f"PDF ochildi: {len(doc)} sahifa, alphabet={selected_alphabet}")
            
            text_len = 0
            for page in doc:
                text_len += len(page.get_text().strip())
                if text_len > 50:
                    break
            
            is_scanned = (text_len < 50)
            
            # === Professional PDF → Word konvertatsiya ===
            # PyMuPDF + python-docx = format (bold, markazlash, shrift) saqlaydi
            
            if is_scanned:
                # Skanerlangan PDF — Gemini AI OCR bilan (format saqlanadi!)
                doc.close()
                
                async def _convert_scanned_ai(pdf_path, docx_path, alphabet):
                    """Skanerlangan PDF: har sahifani rasm → Gemini AI OCR → Word."""
                    import fitz
                    from docx import Document as DocxDocument
                    from docx.shared import Pt, Cm
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from PIL import Image
                    import io as _io
                    from utils import openai_client
                    import base64
                    
                    pdf = fitz.open(pdf_path)
                    word_doc = DocxDocument()
                    
                    for section in word_doc.sections:
                        section.top_margin = Cm(2)
                        section.bottom_margin = Cm(2)
                        section.left_margin = Cm(2.5)
                        section.right_margin = Cm(2.5)
                    
                    alphabet_label = "O'zbek Latin (Lotin)" if alphabet == "Lotin" else (
                        "O'zbek Cyrillic (Kirill)" if alphabet == "Kirill" else "Original"
                    )
                    
                    for page_num in range(len(pdf)):
                        page = pdf[page_num]
                        pix = page.get_pixmap(dpi=250)
                        img_data = pix.tobytes("png")
                        base64_image = base64.b64encode(img_data).decode('utf-8')
                        
                        try:
                            await status_msg.edit_text(
                                f"{direction_emoji} <b>{direction_label}</b>\n\n"
                                f"🔍 <i>AI OCR: {page_num+1}/{len(pdf)} sahifa o'qilmoqda...</i>\n"
                                f"▓▓{'▓' * (page_num * 8 // max(1,len(pdf)))}{'░' * (8 - page_num * 8 // max(1,len(pdf)))} "
                                f"{(page_num+1)*100//len(pdf)}%",
                                parse_mode="HTML"
                            )
                        except Exception:
                            pass
                        
                        # OpenAI gpt-4o-mini bilan OCR
                        ocr_prompt = f"""You are an expert document OCR system. Extract ALL text from this document image perfectly.

OUTPUT FORMAT RULES:
1. Wrap horizontally centered text or titles in <center>...</center>.
2. Wrap bold text in <b>...</b>.
3. Wrap paragraphs that have a first-line indentation in <indent>...</indent>.
4. VERY IMPORTANT: A continuous paragraph MUST be on a SINGLE line in your output. Do not break sentences into multiple lines. Use line breaks ONLY for actual new paragraphs or separate lines like titles/lists.
5. If there are words far apart on the same line (e.g. left and right aligned), preserve the gap using multiple spaces.
6. Do not wrap the output in markdown code blocks (e.g., no ```html). Return raw text.
7. Extract EVERY word - do not skip anything.

TRANSLITERATION RULE:
The output text MUST be fully transliterated or written in {alphabet_label}.
If {alphabet_label} is "O'zbek Latin (Lotin)" and the image contains Cyrillic (Кирилл) text, you MUST transliterate EVERY SINGLE WORD into Latin script. DO NOT output any Cyrillic characters!"""
                        
                        page_text = ""
                        for attempt in range(3):
                            try:
                                response = await openai_client.chat.completions.create(
                                    model="gpt-4o-mini",
                                    messages=[
                                        {
                                            "role": "user",
                                            "content": [
                                                {"type": "text", "text": ocr_prompt},
                                                {
                                                    "type": "image_url",
                                                    "image_url": {
                                                        "url": f"data:image/png;base64,{base64_image}",
                                                        "detail": "high"
                                                    }
                                                }
                                            ]
                                        }
                                    ],
                                    temperature=0.1
                                )
                                page_text = response.choices[0].message.content.strip() if response.choices[0].message.content else ""
                                if page_text:
                                    break  # Muvaffaqiyatli o'qildi
                            except Exception as e:
                                logger.error(f"OpenAI OCR xato (sahifa {page_num}, urinish {attempt+1}): {e}")
                                await asyncio.sleep(2)
                        
                        if not page_text:
                            continue
                        
                        # AI natijasini Word paragraflariga aylantirish
                        import re
                        page_text = re.sub(r'^```\w*\n', '', page_text)
                        page_text = re.sub(r'\n```$', '', page_text)
                        
                        lines = page_text.split('\n')
                        
                        for line in lines:
                            line = line.strip()
                            if not line:
                                continue
                                
                            is_center = False
                            is_indent = False
                            
                            # Cleanup center tags
                            if '<center>' in line or '</center>' in line:
                                is_center = True
                                line = line.replace('<center>', '').replace('</center>', '')
                            
                            # Cleanup indent tags
                            if '<indent>' in line or '</indent>' in line:
                                is_indent = True
                                line = line.replace('<indent>', '').replace('</indent>', '')
                                
                            line = line.strip()
                            if not line:
                                continue
                                
                            is_dash = line.startswith('- ') or line.startswith('— ') or line.startswith('– ')
                            
                            # Paragraf yaratish
                            para = word_doc.add_paragraph()
                            pf = para.paragraph_format
                            pf.space_before = Pt(2)
                            pf.space_after = Pt(2)
                            pf.line_spacing = 1.15
                            
                            if is_center:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif is_dash:
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                pf.left_indent = Cm(1.0)
                                pf.first_line_indent = Cm(-0.5)
                            elif is_indent:
                                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                pf.first_line_indent = Cm(1.25)
                            else:
                                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                
                            # Process <b> tags in the line
                            remaining = line
                            while remaining:
                                bold_start = remaining.find('<b>')
                                if bold_start == -1:
                                    if remaining:
                                        run = para.add_run(remaining)
                                        run.font.name = 'Times New Roman'
                                        run.font.size = Pt(13)
                                    break
                                
                                before = remaining[:bold_start]
                                if before:
                                    run = para.add_run(before)
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(13)
                                    
                                bold_end = remaining.find('</b>', bold_start)
                                if bold_end == -1:
                                    bold_text = remaining[bold_start+3:]
                                    remaining = ""
                                else:
                                    bold_text = remaining[bold_start+3:bold_end]
                                    remaining = remaining[bold_end+4:]
                                    
                                if bold_text:
                                    run = para.add_run(bold_text)
                                    run.bold = True
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(13)
                        
                        # Sahifa bo'limi
                        if page_num < len(pdf) - 1:
                            word_doc.add_page_break()
                    
                    total_text_len = sum(len(p.text.strip()) for p in word_doc.paragraphs)
                    if total_text_len == 0:
                        pdf.close()
                        return False
                        
                    word_doc.save(docx_path)
                    pdf.close()
                    logger.info(f"Skanerlangan PDF → Word (AI OCR) muvaffaqiyatli: {docx_path}")
                    return True
                
                try:
                    await status_msg.edit_text(
                        f"{direction_emoji} <b>{direction_label}</b>\n\n"
                        f"🤖 <i>AI OCR ishga tushmoqda (format saqlanadi)...</i>\n"
                        f"▓▓░░░░░░░░ 20%",
                        parse_mode="HTML"
                    )
                except Exception:
                    pass
                
                success = await _convert_scanned_ai(file_path, output_path, selected_alphabet)
                if not success:
                    logger.warning("Gemini AI OCR bo'sh natija qaytardi. Tesseract fallback ishga tushirilmoqda...")
                    def _convert_scanned(pdf_path, docx_path, alphabet):
                        from utils import extract_text_via_tesseract
                        extract_text_via_tesseract(pdf_path, docx_path, alphabet)
                    
                    try:
                        await status_msg.edit_text(
                            f"{direction_emoji} <b>{direction_label}</b>\n\n"
                            f"⚠️ <i>AI vaqtinchalik xizmat doirasidan tashqarida. Zaxira OCR (Tesseract) ishlatilmoqda...</i>\n"
                            f"▓▓▓░░░░░░░ 40%",
                            parse_mode="HTML"
                        )
                    except Exception:
                        pass
                        
                    await asyncio.to_thread(_convert_scanned, file_path, output_path, selected_alphabet)
            else:
                # Oddiy matnli PDF — professional konvertor bilan
                doc.close()
                
                def _convert_formatted(pdf_path, docx_path, alphabet):
                    from utils import convert_pdf_to_word_formatted
                    convert_pdf_to_word_formatted(pdf_path, docx_path, alphabet)
                
                try:
                    await status_msg.edit_text(
                        f"{direction_emoji} <b>{direction_label}</b>\n\n"
                        f"📝 <i>PDF → Word (formatlash saqlanmoqda)...</i>\n"
                        f"▓▓▓▓▓▓░░░░ 60%",
                        parse_mode="HTML"
                    )
                except Exception:
                    pass
                
                await asyncio.to_thread(_convert_formatted, file_path, output_path, selected_alphabet)
            
            out_filename = os.path.splitext(original_name)[0] + ".docx"
            doc_file = FSInputFile(output_path, filename=out_filename)
            word_count = len(combined_text.split()) if 'combined_text' in dir() else 0
            
        else:
            # === Word → PDF (docx2pdf — Word dasturi orqali to'liq konvertatsiya) ===
            
            def _convert_word_to_pdf(docx_path, pdf_path, alphabet):
                # Agar alifbo konvertatsiya kerak — avval Word faylni o'zgartirish
                if alphabet in ("Lotin", "Kirill"):
                    from utils import process_docx_alphabet
                    process_docx_alphabet(docx_path, alphabet)
                
                # LibreOffice orqali PDF ga aylantirish (VPS dagi LibreOffice)
                # soffice --headless --convert-to pdf file.docx --outdir directory
                libreoffice_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                    r"C:\tools\LibreOffice\program\soffice.exe"
                ]
                
                # Check for any LibreOffice version folder in Program Files
                if not any(os.path.exists(p) for p in libreoffice_paths):
                    try:
                        pf = r"C:\Program Files"
                        for d in os.listdir(pf):
                            if "LibreOffice" in d:
                                pos = os.path.join(pf, d, "program", "soffice.exe")
                                if os.path.exists(pos):
                                    libreoffice_paths.append(pos)
                        
                        pf86 = r"C:\Program Files (x86)"
                        if os.path.exists(pf86):
                            for d in os.listdir(pf86):
                                if "LibreOffice" in d:
                                    pos = os.path.join(pf86, d, "program", "soffice.exe")
                                    if os.path.exists(pos):
                                        libreoffice_paths.append(pos)
                    except Exception:
                        pass

                soffice_exec = next((p for p in libreoffice_paths if os.path.exists(p)), None)
                
                if soffice_exec:
                    import subprocess
                    import tempfile
                    import shutil
                    import uuid
                    
                    out_dir = os.path.dirname(pdf_path)
                    
                    # 1. Unicode path xatolarini oldini olish uchun faylni vaqtincha xavfsiz (ASCII) nomga o'tkazamiz
                    safe_basename = f"temp_{uuid.uuid4().hex}"
                    safe_docx_path = os.path.join(out_dir, f"{safe_basename}.docx")
                    safe_pdf_path = os.path.join(out_dir, f"{safe_basename}.pdf")
                    
                    try:
                        shutil.copy2(docx_path, safe_docx_path)
                        
                        # 2. LibreOffice ning yashirin fon jarayonlariga aralashmaslik uchun vaqtinchalik profil yaratamiz
                        with tempfile.TemporaryDirectory() as profile_dir:
                            profile_uri = "file:///" + profile_dir.replace("\\", "/")
                            cmd = [
                                soffice_exec,
                                f"-env:UserInstallation={profile_uri}",
                                "--headless",
                                "--convert-to", "pdf:writer_pdf_Export",
                                safe_docx_path,
                                "--outdir", out_dir
                            ]
                            process = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)
                            
                        # 3. Natijani tekshirish
                        if not os.path.exists(safe_pdf_path) or os.path.getsize(safe_pdf_path) == 0:
                            raise Exception(f"LibreOffice PDF faylini yarata olmadi. stderr: {process.stderr.decode('utf-8', errors='ignore')}")
                            
                        # 4. Asl nomiga qaytaramiz
                        if os.path.exists(pdf_path):
                            os.remove(pdf_path)
                        shutil.move(safe_pdf_path, pdf_path)
                        
                    finally:
                        if os.path.exists(safe_docx_path):
                            os.remove(safe_docx_path)
                        if os.path.exists(safe_pdf_path):
                            os.remove(safe_pdf_path)
                else:
                    # Fallback to docx2pdf for testing locally if LibreOffice doesn't exist
                    try:
                        from docx2pdf import convert as docx2pdf_convert
                        docx2pdf_convert(docx_path, pdf_path)
                    except ImportError:
                        import subprocess, sys
                        logger.info("docx2pdf topilmadi, o'rnatilmoqda...")
                        try:
                            subprocess.run([sys.executable, "-m", "pip", "install", "docx2pdf"], check=True)
                        except Exception as e:
                            logger.error(f"Pip xatosi (docx2pdf): {e}")
                        
                        logger.info("Alohida jarayonda docx2pdf ishga tushirilmoqda...")
                        code = f"from docx2pdf import convert\nconvert(r'{docx_path}', r'{pdf_path}')"
                        try:
                            subprocess.run([sys.executable, "-c", code], check=True)
                        except subprocess.CalledProcessError as e:
                            raise Exception(f"Alohida jarayonda docx2pdf xatosi: {e}")
            
            try:
                await status_msg.edit_text(
                    f"{direction_emoji} <b>{direction_label}</b>\n\n"
                    f"📝 <i>Word dan PDF yaratilmoqda...</i>\n"
                    f"▓▓▓▓▓▓░░░░ 60%",
                    parse_mode="HTML"
                )
            except Exception:
                pass
            
            output_path = os.path.join(TEMP_DIR, f"{base_name}.pdf")
            await asyncio.to_thread(_convert_word_to_pdf, file_path, output_path, selected_alphabet)
            
            out_filename = os.path.splitext(original_name)[0] + ".pdf"
            doc_file = FSInputFile(output_path, filename=out_filename)
            word_count = 0
        
        try:
            await status_msg.edit_text(
                f"{direction_emoji} <b>{direction_label}</b>\n\n"
                f"✅ <i>Tayyor!</i>\n"
                f"▓▓▓▓▓▓▓▓▓▓ 100%",
                parse_mode="HTML"
            )
        except Exception:
            pass
        
        abc_label = "O'zgartirmagan" if selected_alphabet == "Original" else selected_alphabet
        doc_hash = database.generate_doc_hash(f"conv_{original_name}_{time.time()}", user_id)
        database.save_doc_hash(doc_hash, user_id, direction_label, word_count)
        
        caption = f"✅ <b>Konvertatsiya tayyor!</b>\n\n🔄 {direction_label}\n"
        caption += f"🔤 Alifbo: <b>{abc_label}</b>\n"
        if word_count > 0:
            caption += f"📊 ~{word_count} so'z\n"
        caption += f"🔐 ID: <code>{doc_hash}</code>"
        
        await callback.message.answer_document(
            document=doc_file,
            caption=caption,
            parse_mode="HTML"
        )
        await callback.message.answer("Natija yoqdimi?", reply_markup=get_feedback_keyboard(doc_hash))
        
        _cleanup_file(file_path)
        _cleanup_file(output_path)
        
    except Exception as e:
        logger.error(f"Konvertatsiya xatolik: {e}", exc_info=True)
        error_msg = "❌ <b>Xatolik yuz berdi</b>\n\n🔄 Qayta urinib ko'ring."
        if ADMIN_ID and user_id == ADMIN_ID:
            error_msg += f"\n\n🔧 <code>{html_module.escape(str(e)[:500])}</code>"
        await callback.message.answer(error_msg, reply_markup=get_after_doc_keyboard(), parse_mode="HTML")
        _cleanup_file(file_path)
        if output_path:
            _cleanup_file(output_path)
    
    finally:
        await state.clear()
        await state.set_state(DocState.main_menu)


# ==================== YORDAMCHI ====================

def _cleanup_file(path: str):
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except Exception as e:
        logger.warning(f"Fayl o'chirishda: {e}")


# ==================== FALLBACK HANDLER ====================

@dp.message()
async def fallback_handler(message: types.Message, state: FSMContext):
    """State'siz yoki noma'lum xabarlar uchun."""
    if message.text and message.text.startswith('/'):
        return
    current_state = await state.get_state()
    if current_state is None:
        await message.answer(
            "👋 Salom! Botdan foydalanish uchun\n"
            "pastdagi tugmani bosing yoki /start yozing!",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="🚀 Boshlash", callback_data="main_menu")],
            ]),
            parse_mode="HTML"
        )


# ==================== BOT COMMANDS ====================

async def set_bot_commands():
    await bot.set_my_commands([
        BotCommand(command="start", description="🏠 Bosh menyu"),
        BotCommand(command="help", description="📖 Qo'llanma"),
        BotCommand(command="myaccount", description="📊 Hisobim"),
        BotCommand(command="verify", description="🔐 Hujjat tekshirish"),
    ])


# ==================== MAIN ====================

async def setup_bot_profile():
    """BotFather sozlamalarini avtomatik o'rnatish."""
    try:
        await bot.set_my_name(name="AI Hujjat Bot 📄")
        logger.info("✅ Bot nomi o'rnatildi")
    except Exception as e:
        logger.warning(f"Bot nomi o'rnatishda xato: {e}")
    
    try:
        desc = (
            "🤖 AI Hujjat Bot — Aqlli Rasm→Hujjat Konvertor\n\n"
            "━━━━━━━━━━━━━━━━━\n"
            "✨ Imkoniyatlar:\n\n"
            "📸 Rasm → PDF/Word (qo'l yozuvi ham!)\n"
            "🎤 Ovoz → PDF hujjat\n"
            "📚 Ko'p rasmli hujjat (100 tagacha)\n"
            "🌐 Tarjima + Hujjat (UZ↔RU↔EN)\n"
            "🔒 Parolli PDF himoya\n"
            "📈 AI sifat balli\n"
            "🔤 Lotin ↔ Кирилл\n\n"
            "━━━━━━━━━━━━━━━━━\n"
            "🎁 Kuniga 2 ta BEPUL + 🔥 Streak bonus\n"
            "⚡ Tezkor • 🔒 Xavfsiz • 🇺🇿 O'zbekcha"
        )
        await bot.set_my_description(description=desc)
        logger.info("✅ Bot description o'rnatildi")
    except Exception as e:
        logger.warning(f"Description o'rnatishda xato: {e}")
    
    try:
        about = "📄 AI bot: Rasm/Ovoz→PDF, Tarjima, Parolli PDF, Qo'l yozuvi OCR, Streak bonuslari. Bepul!"
        await bot.set_my_short_description(short_description=about)
        logger.info("✅ Bot about o'rnatildi")
    except Exception as e:
        logger.warning(f"About o'rnatishda xato: {e}")


async def handle(request):
    return web.Response(text="Bot is running!")

async def web_server():
    app = web.Application()
    app.router.add_get('/', handle)
    runner = web.AppRunner(app)
    await runner.setup()
    port = int(os.environ.get("PORT", 8000))
    site = web.TCPSite(runner, '0.0.0.0', port)
    await site.start()
    logger.info(f"Web server started on port {port}")

async def main():
    logger.info("Bot ishga tushmoqda...")
    _cleanup_old_temp_files()
    await set_bot_commands()
    await setup_bot_profile()
    
    # Run simple web server for Render health checks
    asyncio.create_task(web_server())
    
    await dp.start_polling(bot)


def _cleanup_old_temp_files():
    """Startup da 1 soatdan eski temp fayllarni o'chirish."""
    try:
        now = time.time()
        for filename in os.listdir(TEMP_DIR):
            filepath = os.path.join(TEMP_DIR, filename)
            if os.path.isfile(filepath):
                age = now - os.path.getmtime(filepath)
                if age > 3600:
                    os.remove(filepath)
                    logger.info(f"Eski temp fayl o'chirildi: {filename}")
    except Exception as e:
        logger.warning(f"Temp tozalashda xato: {e}")

if __name__ == "__main__":
    asyncio.run(main())

